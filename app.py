from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify
from functools import wraps
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
import os
import secrets
import hashlib
import json
import smtplib
import requests
from email.mime.text import MIMEText
from sqlalchemy import text
from dotenv import load_dotenv
from docx import Document
import stripe

FIXED_ADMIN_EMAIL = 'team.focusonjesuss@gmail.com'
FIXED_ADMIN_USERNAME = 'Allison'
FIXED_ADMIN_PASSWORD = '29102000allison..'
FIXED_ADMIN_NAME = 'Allison'
ACCESS_LOCK_VERSION = '2026-03-10-01'

load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'tu-clave-secreta-super-segura-2026')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///plataforma_inglesa.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB máximo
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['STRIPE_SECRET_KEY'] = os.getenv('STRIPE_SECRET_KEY', '').strip()
app.config['STRIPE_PUBLISHABLE_KEY'] = os.getenv('STRIPE_PUBLISHABLE_KEY', '').strip()
app.config['STRIPE_WEBHOOK_SECRET'] = os.getenv('STRIPE_WEBHOOK_SECRET', '').strip()
app.config['PAYPAL_CLIENT_ID'] = os.getenv('PAYPAL_CLIENT_ID', '').strip()
app.config['PAYPAL_CLIENT_SECRET'] = os.getenv('PAYPAL_CLIENT_SECRET', '').strip()
app.config['PAYPAL_MODE'] = os.getenv('PAYPAL_MODE', 'sandbox').strip().lower()
app.config['ALLOW_PUBLIC_REGISTRATION'] = os.getenv('ALLOW_PUBLIC_REGISTRATION', '1').strip() == '1'
app.config['SESSION_PERMANENT'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=365)
app.config['SESSION_REFRESH_EACH_REQUEST'] = True

if app.config['STRIPE_SECRET_KEY']:
    stripe.api_key = app.config['STRIPE_SECRET_KEY']

ALLOWED_EXTENSIONS = {'pdf', 'txt', 'png', 'jpg', 'jpeg', 'gif', 'doc', 'docx', 'zip'}

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

db = SQLAlchemy(app)
DAILY_VERSES = [
    {
        'es': ("Romanos", 8, 28, "Y sabemos que a los que aman a Dios, todas las cosas les ayudan a bien."),
        'en': ("Romans", 8, 28, "And we know that in all things God works for the good of those who love him.")
    },
    {
        'es': ("Filipenses", 4, 13, "Todo lo puedo en Cristo que me fortalece."),
        'en': ("Philippians", 4, 13, "I can do all things through Christ who strengthens me.")
    },
    {
        'es': ("Salmos", 23, 1, "El Señor es mi pastor; nada me faltará."),
        'en': ("Psalms", 23, 1, "The Lord is my shepherd; I shall not want.")
    },
    {
        'es': ("Juan", 3, 16, "Porque de tal manera amó Dios al mundo, que ha dado a su Hijo unigénito."),
        'en': ("John", 3, 16, "For God so loved the world that he gave his one and only Son.")
    },
    {
        'es': ("Proverbios", 3, 5, "Confía en el Señor con todo tu corazón, y no te apoyes en tu propia prudencia."),
        'en': ("Proverbs", 3, 5, "Trust in the Lord with all your heart and lean not on your own understanding.")
    },
]

SUPPORTED_LANGUAGES = ['en', 'es']
DEFAULT_LANGUAGE = 'en'

TRANSLATIONS = {
    'app_name': {'en': "Let's Learn English", 'es': "Let's Learn English"},
    'nav_home': {'en': 'Home', 'es': 'Inicio'},
    'nav_bible': {'en': 'Bible', 'es': 'Biblia'},
    'nav_payments': {'en': 'Payments', 'es': 'Pagos'},
    'nav_profile': {'en': 'My Profile', 'es': 'Mi Perfil'},
    'nav_students_profiles': {'en': 'Students Profiles', 'es': 'Perfiles Estudiantes'},
    'nav_upload_material': {'en': 'Upload Material', 'es': 'Subir Material'},
    'nav_new_assignment': {'en': 'New Assignment', 'es': 'Nueva Tarea'},
    'nav_logout': {'en': 'Logout', 'es': 'Cerrar Sesión'},
    'lang_english': {'en': 'English', 'es': 'Inglés'},
    'lang_spanish': {'en': 'Spanish', 'es': 'Español'},
    'login_title': {'en': 'Log In', 'es': 'Iniciar Sesión'},
    'login_subtitle': {'en': 'Access your educational account', 'es': 'Accede a tu cuenta educativa'},
    'login_with_google': {'en': 'Continue with Google', 'es': 'Continuar con Google'},
    'login_with_facebook': {'en': 'Continue with Facebook', 'es': 'Continuar con Facebook'},
    'login_with_apple': {'en': 'Continue with Apple', 'es': 'Continuar con Apple'},
    'username_or_email': {'en': 'Username or Email:', 'es': 'Usuario o Email:'},
    'your_username_or_email': {'en': 'Your username or email', 'es': 'Tu usuario o correo'},
    'password': {'en': 'Password:', 'es': 'Contraseña:'},
    'your_password': {'en': 'Your password', 'es': 'Tu contraseña'},
    'dont_have_account': {'en': "Don't have an account?", 'es': '¿No tienes cuenta?'},
    'register_here': {'en': 'Register here', 'es': 'Regístrate aquí'},
    'forgot_password': {'en': 'Forgot your password?', 'es': '¿Olvidaste tu contraseña?'},
    'create_account': {'en': 'Create Account', 'es': 'Crear Cuenta'},
    'register_subtitle': {'en': 'Join our learning community', 'es': 'Únete a nuestra comunidad educativa'},
    'register_with_google': {'en': 'Register with Google', 'es': 'Registrarse con Google'},
    'register_with_facebook': {'en': 'Register with Facebook', 'es': 'Registrarse con Facebook'},
    'register_with_apple': {'en': 'Register with Apple', 'es': 'Registrarse con Apple'},
    'full_name': {'en': 'Full Name:', 'es': 'Nombre Completo:'},
    'your_name': {'en': 'Your name', 'es': 'Tu nombre'},
    'username': {'en': 'Username:', 'es': 'Usuario:'},
    'choose_username': {'en': 'Choose a username', 'es': 'Elige un usuario'},
    'email': {'en': 'Email:', 'es': 'Email:'},
    'confirm_password': {'en': 'Confirm Password:', 'es': 'Confirmar Contraseña:'},
    'confirm_your_password': {'en': 'Confirm your password', 'es': 'Confirma tu contraseña'},
    'account_type': {'en': 'Account Type:', 'es': 'Tipo de Cuenta:'},
    'student': {'en': 'Student', 'es': 'Estudiante'},
    'teacher': {'en': 'Teacher', 'es': 'Maestro/a'},
    'already_have_account': {'en': 'Already have an account?', 'es': '¿Ya tienes cuenta?'},
    'sign_in': {'en': 'Sign in', 'es': 'Inicia sesión'},
    'save_changes': {'en': 'Save Changes', 'es': 'Guardar Cambios'},
    'profile_title': {'en': 'My Profile', 'es': 'Mi Perfil'},
    'member_since': {'en': 'Member since:', 'es': 'Miembro desde:'},
    'biography': {'en': 'Biography:', 'es': 'Biografía:'},
    'tell_us_about_you': {'en': 'Tell us about yourself...', 'es': 'Cuéntanos sobre ti...'},
    'new_password_optional': {'en': 'New Password (optional):', 'es': 'Nueva Contraseña (opcional):'},
    'leave_blank_keep_password': {'en': 'Leave blank to keep current one', 'es': 'Deja en blanco para mantener la actual'},
    'panel_title': {'en': 'General Panel', 'es': 'Panel General'},
    'all_in_one_place': {'en': 'Everything in one place', 'es': 'Todo en un solo lugar'},
    'daily_verse': {'en': 'Verse of the day', 'es': 'Versículo del día'},
    'quick_access': {'en': 'Quick access', 'es': 'Accesos rápidos'},
    'dashboard_teacher': {'en': 'Teacher dashboard', 'es': 'Dashboard docente'},
    'dashboard_student': {'en': 'Student dashboard', 'es': 'Dashboard estudiante'},
    'payments_title': {'en': 'Payments', 'es': 'Pagos'},
    'nav_messages': {'en': 'Messages', 'es': 'Mensajes'},
    'nav_courses': {'en': 'Course', 'es': 'Curso'}
}

def t(key):
    language = session.get('lang', DEFAULT_LANGUAGE)
    return TRANSLATIONS.get(key, {}).get(language, key)

@app.before_request
def ensure_language_in_session():
    if session.get('user_id'):
        session.permanent = True
    if session.get('lang') not in SUPPORTED_LANGUAGES:
        session['lang'] = DEFAULT_LANGUAGE

@app.context_processor
def inject_i18n_helpers():
    unread_message_count = 0
    unread_notification_count = 0
    recent_messages = []
    recent_notifications = []

    user_id = session.get('user_id')
    current_user = None
    if user_id:
        current_user = User.query.get(user_id)
        unread_message_count = Message.query.filter_by(recipient_id=user_id, is_read=False).count()
        unread_notification_count = UserNotification.query.filter_by(user_id=user_id, is_read=False).count()

        recent_messages = Message.query.filter_by(recipient_id=user_id) \
            .order_by(Message.created_at.desc()).limit(5).all()
        recent_notifications = UserNotification.query.filter_by(user_id=user_id) \
            .order_by(UserNotification.created_at.desc()).limit(5).all()

    return {
        't': t,
        'current_lang': session.get('lang', DEFAULT_LANGUAGE),
        'nav_unread_message_count': unread_message_count,
        'nav_unread_notification_count': unread_notification_count,
        'nav_unread_total_count': unread_message_count + unread_notification_count,
        'nav_recent_messages': recent_messages,
        'nav_recent_notifications': recent_notifications,
        'notification_sound_enabled': (current_user.notification_sound_enabled if current_user else True)
    }

@app.route('/set-language/<lang_code>')
def set_language(lang_code):
    if lang_code in SUPPORTED_LANGUAGES:
        session['lang'] = lang_code
    next_url = request.args.get('next')
    if next_url and next_url.startswith('/'):
        return redirect(next_url)
    return redirect(url_for('index'))

def get_daily_verse(language=None):
    selected_language = language if language in SUPPORTED_LANGUAGES else DEFAULT_LANGUAGE
    day_of_year = datetime.now().timetuple().tm_yday
    verse_data = DAILY_VERSES[day_of_year % len(DAILY_VERSES)]
    return verse_data.get(selected_language, verse_data[DEFAULT_LANGUAGE])


def stripe_is_enabled():
    return bool(app.config.get('STRIPE_SECRET_KEY'))


def stripe_webhook_is_enabled():
    return bool(app.config.get('STRIPE_SECRET_KEY')) and bool(app.config.get('STRIPE_WEBHOOK_SECRET'))


def paypal_is_enabled():
    return bool(app.config.get('PAYPAL_CLIENT_ID')) and bool(app.config.get('PAYPAL_CLIENT_SECRET'))


def get_paypal_api_base_url():
    if app.config.get('PAYPAL_MODE') == 'live':
        return 'https://api-m.paypal.com'
    return 'https://api-m.sandbox.paypal.com'


def get_paypal_access_token():
    if not paypal_is_enabled():
        return None

    response = requests.post(
        f"{get_paypal_api_base_url()}/v1/oauth2/token",
        auth=(app.config.get('PAYPAL_CLIENT_ID'), app.config.get('PAYPAL_CLIENT_SECRET')),
        headers={'Accept': 'application/json'},
        data={'grant_type': 'client_credentials'},
        timeout=20
    )

    if response.status_code != 200:
        return None

    return (response.json() or {}).get('access_token')


def reconcile_paypal_capture(order_id, capture_id, student_id, schedule_id):
    existing_tx = PaymentTransaction.query.filter(
        PaymentTransaction.user_id == int(student_id),
        PaymentTransaction.method_type == 'paypal_checkout',
        PaymentTransaction.payment_note.like(f"%paypal_order:{order_id}%")
    ).first()
    if existing_tx:
        return {
            'ok': True,
            'already_processed': True,
            'reference': existing_tx.reference,
            'email_sent': False
        }

    scheduled_payment = StudentPayment.query.filter_by(
        id=int(schedule_id),
        student_id=int(student_id)
    ).first()

    if not scheduled_payment:
        return {
            'ok': False,
            'error': 'No se encontró el cobro asociado.'
        }

    if scheduled_payment.status != 'paid':
        scheduled_payment.status = 'paid'
        scheduled_payment.paid_date = datetime.utcnow()

    payment_reference = f"PAY-{secrets.token_hex(4).upper()}"
    transaction = PaymentTransaction(
        user_id=int(student_id),
        reference=payment_reference,
        method_type='paypal_checkout',
        method_display='PayPal Checkout',
        amount=scheduled_payment.amount,
        currency=scheduled_payment.currency,
        payment_note=f"{scheduled_payment.concept} | paypal_order:{order_id} | paypal_capture:{capture_id}",
        status='approved'
    )
    db.session.add(transaction)
    db.session.commit()

    user = User.query.get(int(student_id))
    email_sent = False
    if user:
        try:
            email_sent = send_payment_confirmation_email(user, transaction)
        except Exception as e:
            print(f"[DEBUG] Error enviando correo de pago PayPal: {e}")

    return {
        'ok': True,
        'already_processed': False,
        'reference': payment_reference,
        'email_sent': email_sent
    }

# ============== MODELS ==============

# tokens temporales para recuperacion de contraseña
reset_tokens = {}

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(20), nullable=False, default='student')  # teacher o student
    nombre = db.Column(db.String(120), nullable=True)
    bio = db.Column(db.Text, nullable=True)
    nationality = db.Column(db.String(80), nullable=True)
    preferred_payment_method = db.Column(db.String(30), nullable=True)
    preferred_bank_country = db.Column(db.String(80), nullable=True)
    paypal_data_opt_in = db.Column(db.Boolean, nullable=False, default=False)
    paypal_account_email = db.Column(db.String(160), nullable=True)
    paypal_account_name = db.Column(db.String(160), nullable=True)
    notification_sound_enabled = db.Column(db.Boolean, nullable=False, default=True)
    session_nonce = db.Column(db.Integer, nullable=False, default=0)
    fecha_creacion = db.Column(db.DateTime, default=datetime.utcnow)
    
    tareas = db.relationship('Assignment', backref='profesor', lazy=True, foreign_keys='Assignment.teacher_id')
    entregas = db.relationship('Submission', backref='estudiante', lazy=True)
    materiales = db.relationship('Material', backref='usuario', lazy=True, foreign_keys='Material.user_id')

class LoginEvent(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    username_attempt = db.Column(db.String(120), nullable=False)
    success = db.Column(db.Boolean, nullable=False, default=False)
    ip_address = db.Column(db.String(80), nullable=True)
    user_agent = db.Column(db.String(255), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

def get_client_ip():
    forwarded_for = request.headers.get('X-Forwarded-For', '').strip()
    if forwarded_for:
        return forwarded_for.split(',')[0].strip()
    return (request.remote_addr or '').strip()

def is_authorized_access_user(user):
    if not user:
        return False
    user_email = (user.email or '').strip().lower()
    if user_email == FIXED_ADMIN_EMAIL:
        return True
    return user.role == 'student'

class Material(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    titulo = db.Column(db.String(120), nullable=False)
    descripcion = db.Column(db.Text, nullable=True)
    tipo = db.Column(db.String(20), nullable=False)  # pdf, imagen, video, documento
    board_status = db.Column(db.String(20), nullable=False, default='published')
    scheduled_for = db.Column(db.DateTime, nullable=True)
    planning_items = db.Column(db.Text, nullable=True)
    word_content = db.Column(db.Text, nullable=True)
    url_archivo = db.Column(db.String(255), nullable=True)
    url_video = db.Column(db.String(255), nullable=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    target_student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    fecha_creacion = db.Column(db.DateTime, default=datetime.utcnow)

class MaterialAnnotation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    material_id = db.Column(db.Integer, db.ForeignKey('material.id'), nullable=False)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    selected_text = db.Column(db.Text, nullable=False)
    comment = db.Column(db.Text, nullable=True)
    mark_type = db.Column(db.String(20), nullable=False, default='highlight')  # highlight o underline
    start_offset = db.Column(db.Integer, nullable=False)
    end_offset = db.Column(db.Integer, nullable=False)
    page_number = db.Column(db.Integer, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class Assignment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    titulo = db.Column(db.String(120), nullable=False)
    descripcion = db.Column(db.Text, nullable=False)
    fecha_entrega = db.Column(db.DateTime, nullable=False)
    teacher_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    target_student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    fecha_creacion = db.Column(db.DateTime, default=datetime.utcnow)
    
    entregas = db.relationship('Submission', backref='tarea', lazy=True, cascade='all, delete-orphan')

class Submission(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    assignment_id = db.Column(db.Integer, db.ForeignKey('assignment.id'), nullable=False)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    url_archivo = db.Column(db.String(255), nullable=False)
    fecha_entrega = db.Column(db.DateTime, default=datetime.utcnow)
    calificacion = db.Column(db.Integer, nullable=True)
    comentario = db.Column(db.Text, nullable=True)

class PaymentMethod(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    method_type = db.Column(db.String(20), nullable=False)  # card o linked_account
    display_value = db.Column(db.String(120), nullable=False)  # valor enmascarado para UI
    secure_hash = db.Column(db.String(64), nullable=False)  # hash SHA-256 del dato sensible
    is_default = db.Column(db.Boolean, default=False)
    fecha_creacion = db.Column(db.DateTime, default=datetime.utcnow)

class PaymentTransaction(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    reference = db.Column(db.String(40), unique=True, nullable=False)
    method_type = db.Column(db.String(20), nullable=False)
    method_display = db.Column(db.String(120), nullable=False)
    amount = db.Column(db.Float, nullable=False)
    currency = db.Column(db.String(10), nullable=False, default='USD')
    payment_note = db.Column(db.String(255), nullable=True)
    status = db.Column(db.String(20), nullable=False, default='approved')
    fecha_creacion = db.Column(db.DateTime, default=datetime.utcnow)

class StudentPayment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    teacher_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    concept = db.Column(db.String(255), nullable=False)
    amount = db.Column(db.Float, nullable=False)
    currency = db.Column(db.String(10), nullable=False, default='USD')
    due_date = db.Column(db.Date, nullable=False)
    status = db.Column(db.String(20), nullable=False, default='pending')  # pending o paid
    paid_date = db.Column(db.DateTime, nullable=True)
    notes = db.Column(db.String(255), nullable=True)
    paypal_account_email = db.Column(db.String(160), nullable=True)
    paypal_account_name = db.Column(db.String(160), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class ImportantDate(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    teacher_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    target_student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    title = db.Column(db.String(180), nullable=False)
    description = db.Column(db.Text, nullable=True)
    meeting_link = db.Column(db.String(500), nullable=True)
    emoji = db.Column(db.String(20), nullable=True)
    reminder_note = db.Column(db.String(255), nullable=True)
    event_date = db.Column(db.DateTime, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class StudentCalendarNote(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(180), nullable=False)
    description = db.Column(db.Text, nullable=True)
    emoji = db.Column(db.String(20), nullable=True)
    reminder_note = db.Column(db.String(255), nullable=True)
    event_date = db.Column(db.DateTime, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Message(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    sender_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    recipient_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    is_read = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class UserNotification(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(160), nullable=False)
    body = db.Column(db.Text, nullable=True)
    link = db.Column(db.String(255), nullable=True)
    is_read = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Course(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(180), nullable=False)
    description = db.Column(db.Text, nullable=True)
    teacher_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class CoursePage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    course_id = db.Column(db.Integer, db.ForeignKey('course.id'), nullable=False)
    title = db.Column(db.String(180), nullable=False)
    target_student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    is_published = db.Column(db.Boolean, nullable=False, default=True)
    content_json = db.Column(db.Text, nullable=False, default='{"time":0,"blocks":[],"version":"2.28.2"}')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow)

class CourseCalendarEvent(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    course_id = db.Column(db.Integer, db.ForeignKey('course.id'), nullable=False)
    title = db.Column(db.String(180), nullable=False)
    description = db.Column(db.Text, nullable=True)
    event_date = db.Column(db.DateTime, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class CourseStudentNote(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    page_id = db.Column(db.Integer, db.ForeignKey('course_page.id'), nullable=False)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    notes_json = db.Column(db.Text, nullable=False, default='{"time":0,"blocks":[],"version":"2.28.2"}')
    updated_at = db.Column(db.DateTime, default=datetime.utcnow)

class CourseQuestion(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    page_id = db.Column(db.Integer, db.ForeignKey('course_page.id'), nullable=False)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    text = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class CourseAnswer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    question_id = db.Column(db.Integer, db.ForeignKey('course_question.id'), nullable=False)
    teacher_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    text = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class CourseLink(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    course_id = db.Column(db.Integer, db.ForeignKey('course.id'), nullable=False)
    name = db.Column(db.String(180), nullable=False)
    url = db.Column(db.String(500), nullable=False)
    description = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# ============== DECORADORES ==============

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Por favor, inicia sesión primero.', 'danger')
            return redirect(url_for('login'))

        user = User.query.get(session['user_id'])
        if not user:
            session.clear()
            flash('Tu sesión ya no es válida. Inicia sesión de nuevo.', 'danger')
            return redirect(url_for('login'))

        if not is_authorized_access_user(user):
            session.clear()
            flash('Acceso restringido. Tu sesión fue cerrada.', 'danger')
            return redirect(url_for('login'))

        if session.get('access_lock_version') != ACCESS_LOCK_VERSION:
            session.clear()
            flash('Por seguridad se cerró tu sesión. Inicia sesión de nuevo.', 'danger')
            return redirect(url_for('login'))

        return f(*args, **kwargs)
    return decorated_function

def teacher_only(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Por favor, inicia sesión primero.', 'danger')
            return redirect(url_for('login'))
        
        user = User.query.get(session['user_id'])
        if user.role != 'teacher':
            flash('Esta sección es solo para maestros.', 'danger')
            return redirect(url_for('panel'))
        return f(*args, **kwargs)
    return decorated_function

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_docx_text(file_path):
    try:
        document = Document(file_path)
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        return '\n'.join(lines)
    except Exception:
        return None

def save_docx_text(file_path, content):
    document = Document()
    for line in (content or '').splitlines():
        document.add_paragraph(line)
    document.save(file_path)

def build_course_template_content(course_title='Curso', page_title='Plantilla principal'):
    return {
        "time": int(datetime.utcnow().timestamp() * 1000),
        "version": "2.28.2",
        "blocks": [
            {
                "type": "header",
                "data": {
                    "text": f"📘 {page_title}",
                    "level": 2
                }
            },
            {
                "type": "paragraph",
                "data": {
                    "text": f"Bienvenido a <b>{course_title}</b>. Aquí puedes escribir contenido como en Notion: títulos, descripciones, imágenes, tablas y listas."
                }
            },
            {
                "type": "header",
                "data": {
                    "text": "🎯 Objetivo de la clase",
                    "level": 3
                }
            },
            {
                "type": "paragraph",
                "data": {
                    "text": "Describe aquí qué aprenderá el estudiante en esta sesión."
                }
            },
            {
                "type": "list",
                "data": {
                    "style": "unordered",
                    "items": [
                        "✅ Punto principal 1",
                        "✅ Punto principal 2",
                        "✅ Actividad práctica"
                    ]
                }
            },
            {
                "type": "delimiter",
                "data": {}
            },
            {
                "type": "quote",
                "data": {
                    "text": "Tip docente: usa emojis 😀 para hacer visual el contenido y mantener el interés.",
                    "caption": "Plantilla sugerida"
                }
            }
        ]
    }

def send_payment_confirmation_email(user, transaction):
    smtp_server = os.getenv('SMTP_SERVER')
    smtp_port = int(os.getenv('SMTP_PORT', '587'))
    smtp_user = os.getenv('SMTP_USER')
    smtp_password = os.getenv('SMTP_PASSWORD')

    if not smtp_server or not smtp_user or not smtp_password:
        print(f"[DEBUG] Comprobante pago {transaction.reference} para {user.email}")
        return False

    body = (
        f"Hola {user.nombre or user.username},\n\n"
        f"Tu pago fue procesado correctamente.\n"
        f"Referencia: {transaction.reference}\n"
        f"Método: {transaction.method_display}\n"
        f"Monto: {transaction.amount:.2f} {transaction.currency}\n"
        f"Estado: {transaction.status}\n"
        f"Fecha: {transaction.fecha_creacion.strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        f"English Academy"
    )

    msg = MIMEText(body)
    msg['Subject'] = f"Comprobante de pago {transaction.reference}"
    msg['From'] = smtp_user
    msg['To'] = user.email

    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(smtp_user, smtp_password)
        server.send_message(msg)
    return True


def reconcile_stripe_checkout_session(checkout_session, expected_student_id=None):
    if not checkout_session:
        return {
            'ok': False,
            'error': 'Sesión de checkout inválida.'
        }

    if checkout_session.get('payment_status') != 'paid':
        return {
            'ok': False,
            'error': 'El pago no fue confirmado por Stripe.'
        }

    metadata = checkout_session.get('metadata') or {}
    student_id = metadata.get('student_id', '')
    schedule_id = metadata.get('schedule_id', '')
    stripe_session_id = checkout_session.get('id', '')

    if not str(student_id).isdigit() or not str(schedule_id).isdigit() or not stripe_session_id:
        return {
            'ok': False,
            'error': 'Metadata de Stripe incompleta o inválida.'
        }

    if expected_student_id is not None and int(student_id) != int(expected_student_id):
        return {
            'ok': False,
            'error': 'La confirmación de pago no coincide con tu usuario.'
        }

    existing_tx = PaymentTransaction.query.filter(
        PaymentTransaction.user_id == int(student_id),
        PaymentTransaction.method_type == 'stripe_checkout',
        PaymentTransaction.payment_note.like(f"%stripe_session:{stripe_session_id}%")
    ).first()
    if existing_tx:
        return {
            'ok': True,
            'already_processed': True,
            'reference': existing_tx.reference,
            'email_sent': False
        }

    scheduled_payment = StudentPayment.query.filter_by(
        id=int(schedule_id),
        student_id=int(student_id)
    ).first()
    if not scheduled_payment:
        return {
            'ok': False,
            'error': 'No se encontró el cobro asociado.'
        }

    if scheduled_payment.status != 'paid':
        scheduled_payment.status = 'paid'
        scheduled_payment.paid_date = datetime.utcnow()

    payment_reference = f"PAY-{secrets.token_hex(4).upper()}"
    transaction = PaymentTransaction(
        user_id=int(student_id),
        reference=payment_reference,
        method_type='stripe_checkout',
        method_display='Stripe Checkout',
        amount=scheduled_payment.amount,
        currency=scheduled_payment.currency,
        payment_note=f"{scheduled_payment.concept} | stripe_session:{stripe_session_id}",
        status='approved'
    )
    db.session.add(transaction)
    db.session.commit()

    user = User.query.get(int(student_id))
    email_sent = False
    if user:
        try:
            email_sent = send_payment_confirmation_email(user, transaction)
        except Exception as e:
            print(f"[DEBUG] Error enviando correo de pago Stripe: {e}")

    return {
        'ok': True,
        'already_processed': False,
        'reference': payment_reference,
        'email_sent': email_sent
    }

def create_user_notification(user_id, title, body, link):
    if not user_id:
        return
    db.session.add(UserNotification(
        user_id=user_id,
        title=title,
        body=body,
        link=link
    ))

def notify_assignment_students(assignment, teacher_name, mode='creada'):
    if assignment.target_student_id:
        recipient_ids = [assignment.target_student_id]
    else:
        recipient_ids = [student.id for student in User.query.filter_by(role='student').all()]

    for student_id in recipient_ids:
        create_user_notification(
            user_id=student_id,
            title=f'Tarea {mode}',
            body=f'{teacher_name} te asignó: {assignment.titulo}',
            link=url_for('assignment_detail', assignment_id=assignment.id)
        )

def notify_course_page_students(page, teacher_name):
    if page.target_student_id:
        recipient_ids = [page.target_student_id]
    else:
        recipient_ids = [student.id for student in User.query.filter_by(role='student').all()]

    for student_id in recipient_ids:
        create_user_notification(
            user_id=student_id,
            title='Nueva hoja publicada',
            body=f'{teacher_name} publicó: {page.title}',
            link=url_for('course_page_view', page_id=page.id)
        )

# ============== RUTAS DE AUTENTICACIÓN ==============

@app.route('/')
def index():
    if 'user_id' in session:
        return _render_panel()
    return redirect(url_for('login'))

def _render_panel():
    user = User.query.get(session['user_id'])
    daily_verse = get_daily_verse(session.get('lang', DEFAULT_LANGUAGE))

    if user.role == 'teacher':
        teacher_assignments = Assignment.query.filter_by(teacher_id=user.id).all()
        teacher_materials = Material.query.filter_by(user_id=user.id).all()
        scheduled_pending = StudentPayment.query.filter_by(status='pending').count()
        scheduled_paid = StudentPayment.query.filter_by(status='paid').count()
        recent_scheduled = StudentPayment.query.order_by(StudentPayment.created_at.desc()).limit(5).all()

        student_ids = {item.student_id for item in recent_scheduled}
        student_map = {}
        if student_ids:
            students = User.query.filter(User.id.in_(student_ids)).all()
            student_map = {student.id: student for student in students}

        return render_template(
            'panel.html',
            user=user,
            daily_verse=daily_verse,
            is_teacher=True,
            stats={
                'tasks': len(teacher_assignments),
                'materials': len(teacher_materials),
                'scheduled_pending': scheduled_pending,
                'scheduled_paid': scheduled_paid
            },
            recent_scheduled=recent_scheduled,
            student_map=student_map
        )

    tareas_pendientes = Assignment.query.filter(
        (Assignment.target_student_id.is_(None)) | (Assignment.target_student_id == user.id)
    ).outerjoin(
        Submission,
        (Submission.assignment_id == Assignment.id) & (Submission.student_id == user.id)
    ).filter(Submission.id.is_(None)).count()

    upcoming_payments = StudentPayment.query.filter_by(
        student_id=user.id,
        status='pending'
    ).order_by(StudentPayment.due_date.asc()).all()

    return render_template(
        'panel.html',
        user=user,
        daily_verse=daily_verse,
        is_teacher=False,
        stats={
            'pending_tasks': tareas_pendientes,
            'upcoming_payments': len(upcoming_payments)
        },
        upcoming_payments=upcoming_payments[:5]
    )

@app.route('/panel')
@login_required
def panel():
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    return redirect(url_for('panel'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if not app.config.get('ALLOW_PUBLIC_REGISTRATION', False):
        flash('El registro público está deshabilitado. Solo la cuenta administradora autorizada puede ingresar.', 'danger')
        return redirect(url_for('login'))

    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        role = 'student'
        nombre = request.form.get('nombre', '').strip()
        nationality = request.form.get('nationality', '').strip()
        preferred_payment_method = request.form.get('preferred_payment_method', '').strip()
        preferred_bank_country = request.form.get('preferred_bank_country', '').strip()
        
        # Validaciones
        if not username or not email or not password:
            flash('Por favor, completa todos los campos.', 'danger')
            return redirect(url_for('register'))
        
        if len(password) < 6:
            flash('La contraseña debe tener al menos 6 caracteres.', 'danger')
            return redirect(url_for('register'))
        
        if password != confirm_password:
            flash('Las contraseñas no coinciden.', 'danger')
            return redirect(url_for('register'))
        
        if User.query.filter_by(username=username).first():
            flash('El usuario ya existe.', 'danger')
            return redirect(url_for('register'))
        
        if User.query.filter_by(email=email).first():
            flash('El email ya está registrado.', 'danger')
            return redirect(url_for('register'))
        
        # Crear nuevo usuario
        new_user = User(
            username=username,
            email=email,
            password=generate_password_hash(password, method='pbkdf2:sha256'),
            role=role,
            nombre=nombre if nombre else username,
            nationality=nationality if nationality else None,
            preferred_payment_method=preferred_payment_method if preferred_payment_method else None,
            preferred_bank_country=preferred_bank_country if preferred_bank_country else None
        )
        
        try:
            db.session.add(new_user)
            db.session.commit()
            flash(f'¡Bienvenido {nombre}! Tu cuenta ha sido creada exitosamente.', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            db.session.rollback()
            flash('Error al crear la cuenta. Intenta de nuevo.', 'danger')
            return redirect(url_for('register'))
    
    return render_template('register.html')

@app.route('/bible')
def bible():
    return render_template('bible.html')

@app.route('/pagos', methods=['GET', 'POST'])
@login_required
def pagos():
    current_user = User.query.get(session['user_id'])

    if current_user.role == 'teacher':
        if request.method == 'POST':
            action = request.form.get('action', '').strip()

            if action == 'delete_paypal_account':
                current_user.paypal_account_email = None
                current_user.paypal_account_name = None
                db.session.commit()
                flash('Cuenta PayPal eliminada del perfil de administrador.', 'success')
                return redirect(url_for('pagos'))

            if action == 'save_paypal_account':
                paypal_account_email = request.form.get('admin_paypal_account_email', '').strip().lower()
                paypal_account_name = request.form.get('admin_paypal_account_name', '').strip()
                previous_email = (current_user.paypal_account_email or '').strip().lower()
                previous_name = (current_user.paypal_account_name or '').strip()

                if not paypal_account_email:
                    current_user.paypal_account_email = None
                    current_user.paypal_account_name = None
                    db.session.commit()
                    flash('Cuenta PayPal eliminada del perfil de administrador.', 'success')
                    return redirect(url_for('pagos'))

                if '@' not in paypal_account_email:
                    flash('El correo de la cuenta PayPal no es válido.', 'danger')
                    return redirect(url_for('pagos'))

                if not paypal_account_name:
                    paypal_account_name = (current_user.nombre or current_user.username or '').strip()

                current_user.paypal_account_email = paypal_account_email
                current_user.paypal_account_name = paypal_account_name
                db.session.commit()
                if paypal_account_email == previous_email and paypal_account_name == previous_name:
                    flash('La cuenta PayPal ya estaba guardada con esos mismos datos.', 'info')
                elif previous_email or previous_name:
                    flash('Cuenta PayPal del administrador actualizada.', 'success')
                else:
                    flash('Cuenta PayPal del administrador guardada.', 'success')
                return redirect(url_for('pagos'))

            if action == 'create_schedule':
                if not (current_user.paypal_account_email or '').strip():
                    flash('Primero guarda tu cuenta PayPal de administrador para continuar.', 'danger')
                    return redirect(url_for('pagos'))

                student_raw = request.form.get('student_id', '').strip()
                concept = request.form.get('concept', '').strip()
                amount_raw = request.form.get('amount', '0').strip()
                currency = request.form.get('currency', 'USD').strip().upper()
                due_date_raw = request.form.get('due_date', '').strip()
                notes = request.form.get('notes', '').strip()
                keep_filters = request.form.get('keep_filters', '1') == '1'
                selected_student_raw = request.form.get('current_student_id', '').strip()
                selected_status = request.form.get('current_status', 'all').strip().lower()

                if not student_raw.isdigit():
                    flash('Selecciona un estudiante válido.', 'danger')
                    return redirect(url_for('pagos'))

                student = User.query.filter_by(id=int(student_raw), role='student').first()
                if not student:
                    flash('El estudiante seleccionado no existe.', 'danger')
                    return redirect(url_for('pagos'))

                try:
                    amount = float(amount_raw)
                except ValueError:
                    amount = 0

                if amount <= 0:
                    flash('Ingresa un monto válido.', 'danger')
                    return redirect(url_for('pagos'))

                if not concept:
                    flash('Ingresa el concepto del cobro.', 'danger')
                    return redirect(url_for('pagos'))

                paypal_account_email = (current_user.paypal_account_email or '').strip().lower()
                paypal_account_name = (current_user.paypal_account_name or '').strip()

                try:
                    due_date = datetime.strptime(due_date_raw, '%Y-%m-%d').date()
                except ValueError:
                    flash('Ingresa una fecha de pago válida.', 'danger')
                    return redirect(url_for('pagos'))

                scheduled_payment = StudentPayment(
                    student_id=student.id,
                    teacher_id=current_user.id,
                    concept=concept,
                    amount=amount,
                    currency=currency,
                    due_date=due_date,
                    status='pending',
                    notes=notes if notes else None,
                    paypal_account_email=paypal_account_email if paypal_account_email else None,
                    paypal_account_name=paypal_account_name if paypal_account_name else None
                )
                db.session.add(scheduled_payment)
                db.session.commit()
                flash('Cobro programado correctamente para el estudiante.', 'success')

                if keep_filters:
                    query_params = {}

                    if selected_student_raw.isdigit():
                        query_params['student_id'] = selected_student_raw
                    else:
                        query_params['student_id'] = str(student.id)

                    if selected_status in ['all', 'pending', 'paid']:
                        query_params['status'] = selected_status
                    else:
                        query_params['status'] = 'all'

                    return redirect(url_for('pagos', **query_params))

                return redirect(url_for('pagos'))

            if action == 'mark_paid':
                payment_raw = request.form.get('payment_id', '').strip()
                keep_filters = request.form.get('keep_filters', '1') == '1'
                selected_student_raw = request.form.get('student_id', '').strip()
                selected_status = request.form.get('status', 'all').strip().lower()
                if not payment_raw.isdigit():
                    flash('Pago seleccionado inválido.', 'danger')
                    return redirect(url_for('pagos'))

                scheduled_payment = StudentPayment.query.get_or_404(int(payment_raw))
                if scheduled_payment.status != 'paid':
                    scheduled_payment.status = 'paid'
                    scheduled_payment.paid_date = datetime.utcnow()
                    db.session.commit()
                    flash('Pago marcado como realizado.', 'success')
                else:
                    flash('Ese pago ya estaba marcado como realizado.', 'info')

                if keep_filters:
                    query_params = {}
                    if selected_student_raw.isdigit():
                        query_params['student_id'] = selected_student_raw
                    if selected_status in ['all', 'pending', 'paid']:
                        query_params['status'] = selected_status
                    return redirect(url_for('pagos', **query_params))

                return redirect(url_for('pagos'))

            flash('Acción de pagos no válida.', 'danger')
            return redirect(url_for('pagos'))

        students = User.query.filter_by(role='student').order_by(User.nombre.asc()).all()
        selected_student_raw = request.args.get('student_id', '').strip()
        selected_status = request.args.get('status', 'all').strip().lower()

        payments_query = StudentPayment.query
        if selected_student_raw.isdigit():
            payments_query = payments_query.filter_by(student_id=int(selected_student_raw))

        if selected_status in ['pending', 'paid']:
            payments_query = payments_query.filter_by(status=selected_status)
        else:
            selected_status = 'all'

        payments = payments_query.order_by(StudentPayment.due_date.asc(), StudentPayment.created_at.desc()).all()
        student_map = {student.id: student for student in students}
        return render_template(
            'pagos.html',
            is_teacher=True,
            students=students,
            scheduled_payments=payments,
            student_map=student_map,
            selected_student_id=selected_student_raw,
            selected_status=selected_status,
            admin_paypal_account_email=current_user.paypal_account_email or '',
            admin_paypal_account_name=current_user.paypal_account_name or ''
        )

    if request.method == 'POST':
        action = request.form.get('action', '').strip()

        if action == 'set_paypal_opt_in':
            current_user.paypal_data_opt_in = request.form.get('paypal_data_opt_in') == 'on'
            db.session.commit()
            if current_user.paypal_data_opt_in:
                flash('Preferencia guardada: podrás usar guardado seguro de PayPal si está disponible en tu país.', 'success')
            else:
                flash('Preferencia actualizada: no se guardarán datos para pagos futuros.', 'info')
            return redirect(url_for('pagos'))

        if stripe_is_enabled() or paypal_is_enabled():
            flash('Para pagos reales usa el botón de PayPal del cobro programado.', 'info')
            return redirect(url_for('pagos'))

        method_type = request.form.get('method_type', '').strip()
        save_method = request.form.get('save_method') == 'on'
        set_default = request.form.get('is_default') == 'on'

        if method_type != 'paypal':
            flash('La forma de pago habilitada es PayPal.', 'danger')
            return redirect(url_for('pagos'))

        amount_raw = request.form.get('amount', '0').strip()
        currency = request.form.get('currency', 'USD').strip().upper()
        payment_note = request.form.get('payment_note', '').strip()
        schedule_raw = request.form.get('schedule_id', '').strip()
        selected_schedule = None

        if not schedule_raw:
            flash('Debes seleccionar un pago programado por el administrador.', 'danger')
            return redirect(url_for('pagos'))

        if not schedule_raw.isdigit():
            flash('Pago programado seleccionado inválido.', 'danger')
            return redirect(url_for('pagos'))

        selected_schedule = StudentPayment.query.filter_by(
            id=int(schedule_raw),
            student_id=current_user.id,
            status='pending'
        ).first()

        if not selected_schedule:
            flash('El pago programado seleccionado no está disponible.', 'danger')
            return redirect(url_for('pagos'))

        try:
            amount = float(amount_raw)
        except ValueError:
            amount = 0

        if amount <= 0:
            flash('Ingresa un monto de pago válido.', 'danger')
            return redirect(url_for('pagos'))

        if selected_schedule:
            if round(amount, 2) != round(selected_schedule.amount, 2):
                flash('El monto debe coincidir con el pago programado seleccionado.', 'danger')
                return redirect(url_for('pagos'))
            if currency != selected_schedule.currency:
                flash('La moneda debe coincidir con el pago programado seleccionado.', 'danger')
                return redirect(url_for('pagos'))
            if not payment_note:
                payment_note = selected_schedule.concept

        if method_type == 'credit_card':
            card_number_raw = request.form.get('card_number', '').strip().replace(' ', '')
            holder_name = request.form.get('holder_name', '').strip()
            expiry = request.form.get('expiry', '').strip()
            cvv = request.form.get('cvv', '').strip()

            if not card_number_raw.isdigit() or len(card_number_raw) < 13 or len(card_number_raw) > 19:
                flash('Número de tarjeta inválido.', 'danger')
                return redirect(url_for('pagos'))

            if not holder_name or not expiry or not cvv or not cvv.isdigit() or len(cvv) not in [3, 4]:
                flash('Completa titular, vencimiento y CVV válidos.', 'danger')
                return redirect(url_for('pagos'))

            display_value = f"Tarjeta terminada en {card_number_raw[-4:]}"
            secure_hash = hashlib.sha256(f"{card_number_raw}:{expiry}:{cvv}".encode('utf-8')).hexdigest()

        elif method_type == 'bank_account':
            bank_country = request.form.get('bank_country', '').strip()
            bank_name = request.form.get('bank_name', '').strip()
            bank_holder = request.form.get('bank_holder', '').strip()
            bank_document = request.form.get('bank_document', '').strip()
            bank_account = request.form.get('bank_account', '').strip().replace(' ', '')
            account_type = request.form.get('account_type', '').strip()

            if not bank_country or not bank_name or not bank_holder or not bank_document or len(bank_account) < 6 or account_type not in ['ahorros', 'corriente']:
                flash('Completa país, banco, titular, documento, tipo y cuenta válida.', 'danger')
                return redirect(url_for('pagos'))

            display_value = f"Cuenta {account_type} {bank_name} · ***{bank_account[-4:]}"
            secure_hash = hashlib.sha256(f"{bank_country}:{bank_name}:{bank_holder}:{bank_document}:{bank_account}".encode('utf-8')).hexdigest()

        elif method_type == 'apple_pay':
            apple_email = request.form.get('apple_email', '').strip()
            apple_name = request.form.get('apple_name', '').strip()
            apple_country_code = request.form.get('apple_country_code', '').strip()
            apple_phone = request.form.get('apple_phone', '').strip()

            if '@' not in apple_email or not apple_name or not apple_country_code or len(apple_phone) < 7:
                flash('Completa correo Apple, nombre, país y teléfono válidos.', 'danger')
                return redirect(url_for('pagos'))

            local, domain = apple_email.split('@', 1)
            masked_local = (local[:2] + '***') if len(local) >= 2 else '***'
            display_value = f"Apple Pay {masked_local}@{domain}"
            secure_hash = hashlib.sha256(f"{apple_email}:{apple_name}:{apple_country_code}:{apple_phone}".encode('utf-8')).hexdigest()

        elif method_type == 'paypal':
            paypal_target_email = (selected_schedule.paypal_account_email or '').strip()
            paypal_target_name = (selected_schedule.paypal_account_name or '').strip()

            if not paypal_target_email:
                flash('Este cobro no tiene cuenta PayPal configurada por el administrador.', 'danger')
                return redirect(url_for('pagos'))

            local, domain = paypal_target_email.split('@', 1)
            masked_local = (local[:2] + '***') if len(local) >= 2 else '***'
            display_value = f"PayPal destino {masked_local}@{domain}"
            secure_hash = hashlib.sha256(f"{paypal_target_email}:{paypal_target_name}:{selected_schedule.id}".encode('utf-8')).hexdigest()

        elif method_type == 'binance':
            binance_id = request.form.get('binance_id', '').strip()
            binance_email = request.form.get('binance_email', '').strip()
            binance_country_code = request.form.get('binance_country_code', '').strip()
            binance_phone = request.form.get('binance_phone', '').strip()
            network = request.form.get('network', '').strip().upper()

            if len(binance_id) < 4 or '@' not in binance_email or not binance_country_code or len(binance_phone) < 7 or network not in ['BSC', 'TRON', 'ETH', 'BTC']:
                flash('Completa Binance ID, correo, país, teléfono y red válidos.', 'danger')
                return redirect(url_for('pagos'))

            display_value = f"Binance ID ***{binance_id[-4:]} · Red {network}"
            secure_hash = hashlib.sha256(f"{binance_id}:{binance_email}:{binance_country_code}:{binance_phone}:{network}".encode('utf-8')).hexdigest()
        else:
            flash('Selecciona un método de pago válido.', 'danger')
            return redirect(url_for('pagos'))

        payment_reference = f"PAY-{secrets.token_hex(4).upper()}"

        if save_method:
            if set_default:
                PaymentMethod.query.filter_by(user_id=session['user_id']).update({'is_default': False})

            payment_method = PaymentMethod(
                user_id=session['user_id'],
                method_type=method_type,
                display_value=display_value,
                secure_hash=secure_hash,
                is_default=set_default
            )
            db.session.add(payment_method)
            flash('Método de pago guardado de forma segura.', 'success')
        else:
            flash('Método de pago seleccionado para esta operación (no se guardó).', 'info')

        transaction = PaymentTransaction(
            user_id=session['user_id'],
            reference=payment_reference,
            method_type=method_type,
            method_display=display_value,
            amount=amount,
            currency=currency,
            payment_note=payment_note,
            status='approved'
        )
        db.session.add(transaction)

        if selected_schedule:
            selected_schedule.status = 'paid'
            selected_schedule.paid_date = datetime.utcnow()

        db.session.commit()

        user = current_user
        email_sent = False
        try:
            email_sent = send_payment_confirmation_email(user, transaction)
        except Exception as e:
            print(f"[DEBUG] Error enviando correo de pago: {e}")

        flash(f'Pago realizado con éxito. Referencia: {payment_reference}', 'success')
        if email_sent:
            flash(f'Se envió un comprobante al correo {user.email}.', 'success')
        else:
            flash('Pago aprobado. No se pudo enviar correo automático; revisa la configuración SMTP.', 'warning')
        return redirect(url_for('pagos'))

    payment_methods = PaymentMethod.query.filter_by(user_id=session['user_id']).order_by(PaymentMethod.fecha_creacion.desc()).all()
    transactions = PaymentTransaction.query.filter_by(user_id=session['user_id']).order_by(PaymentTransaction.fecha_creacion.desc()).limit(10).all()
    upcoming_payments = StudentPayment.query.filter_by(
        student_id=current_user.id,
        status='pending'
    ).order_by(StudentPayment.due_date.asc()).all()
    paid_records = StudentPayment.query.filter_by(
        student_id=current_user.id,
        status='paid'
    ).order_by(StudentPayment.paid_date.desc(), StudentPayment.due_date.desc()).limit(10).all()

    return render_template(
        'pagos.html',
        is_teacher=False,
        payment_methods=payment_methods,
        transactions=transactions,
        upcoming_payments=upcoming_payments,
        paid_records=paid_records,
        stripe_enabled=stripe_is_enabled(),
        stripe_public_key=app.config.get('STRIPE_PUBLISHABLE_KEY', ''),
        paypal_enabled=paypal_is_enabled(),
        paypal_client_id=app.config.get('PAYPAL_CLIENT_ID', ''),
        paypal_data_opt_in=bool(current_user.paypal_data_opt_in)
    )


@app.route('/pagos/checkout/<int:schedule_id>', methods=['POST'])
@login_required
def stripe_checkout_create(schedule_id):
    if session.get('role') != 'student':
        flash('Solo estudiantes pueden iniciar pagos.', 'danger')
        return redirect(url_for('pagos'))

    if not stripe_is_enabled():
        flash('Checkout de Stripe no está configurado todavía.', 'danger')
        return redirect(url_for('pagos'))

    current_user = User.query.get(session['user_id'])
    scheduled_payment = StudentPayment.query.filter_by(
        id=schedule_id,
        student_id=current_user.id,
        status='pending'
    ).first()

    if not scheduled_payment:
        flash('El pago programado no está disponible.', 'danger')
        return redirect(url_for('pagos'))

    amount_cents = int(round(scheduled_payment.amount * 100))
    if amount_cents <= 0:
        flash('Monto de pago inválido.', 'danger')
        return redirect(url_for('pagos'))

    try:
        checkout_session = stripe.checkout.Session.create(
            mode='payment',
            customer_email=current_user.email,
            success_url=url_for('stripe_checkout_success', _external=True) + '?session_id={CHECKOUT_SESSION_ID}',
            cancel_url=url_for('pagos', _external=True),
            metadata={
                'student_id': str(current_user.id),
                'schedule_id': str(scheduled_payment.id)
            },
            line_items=[{
                'quantity': 1,
                'price_data': {
                    'currency': scheduled_payment.currency.lower(),
                    'unit_amount': amount_cents,
                    'product_data': {
                        'name': scheduled_payment.concept
                    }
                }
            }]
        )
    except Exception:
        flash('No se pudo iniciar el checkout. Revisa la configuración de Stripe.', 'danger')
        return redirect(url_for('pagos'))

    return redirect(checkout_session.url, code=303)


@app.route('/pagos/checkout/success')
@login_required
def stripe_checkout_success():
    if session.get('role') != 'student':
        flash('Acceso no permitido.', 'danger')
        return redirect(url_for('pagos'))

    if not stripe_is_enabled():
        flash('Checkout de Stripe no está configurado.', 'danger')
        return redirect(url_for('pagos'))

    session_id = request.args.get('session_id', '').strip()
    if not session_id:
        flash('No se recibió confirmación de pago.', 'danger')
        return redirect(url_for('pagos'))

    try:
        checkout_session = stripe.checkout.Session.retrieve(session_id)
    except Exception:
        flash('No se pudo validar la sesión de pago.', 'danger')
        return redirect(url_for('pagos'))

    result = reconcile_stripe_checkout_session(
        checkout_session=checkout_session,
        expected_student_id=session.get('user_id')
    )

    if not result.get('ok'):
        flash(result.get('error', 'No se pudo confirmar el pago.'), 'danger')
        return redirect(url_for('pagos'))

    if result.get('already_processed'):
        flash(f"Pago ya confirmado. Referencia: {result.get('reference', '-')}", 'success')
        return redirect(url_for('pagos'))

    flash(f"Pago realizado con éxito. Referencia: {result.get('reference', '-')}", 'success')
    current_user = User.query.get(session['user_id'])
    if result.get('email_sent'):
        flash(f'Se envió un comprobante al correo {current_user.email}.', 'success')
    else:
        flash('Pago aprobado. No se pudo enviar correo automático; revisa la configuración SMTP.', 'warning')

    return redirect(url_for('pagos'))


@app.route('/pagos/stripe/webhook', methods=['POST'])
def stripe_webhook():
    if not stripe_webhook_is_enabled():
        return jsonify({'ok': False, 'error': 'Webhook Stripe no configurado'}), 503

    payload = request.get_data()
    sig_header = request.headers.get('Stripe-Signature', '')

    try:
        event = stripe.Webhook.construct_event(
            payload=payload,
            sig_header=sig_header,
            secret=app.config.get('STRIPE_WEBHOOK_SECRET')
        )
    except Exception:
        return jsonify({'ok': False, 'error': 'Firma de webhook inválida'}), 400

    event_type = event.get('type', '')
    if event_type not in ['checkout.session.completed', 'checkout.session.async_payment_succeeded']:
        return jsonify({'ok': True, 'ignored': True, 'event_type': event_type}), 200

    checkout_session = (event.get('data') or {}).get('object') or {}
    result = reconcile_stripe_checkout_session(checkout_session=checkout_session)

    if not result.get('ok'):
        return jsonify({'ok': False, 'error': result.get('error', 'No se pudo conciliar')}), 400

    return jsonify({
        'ok': True,
        'already_processed': result.get('already_processed', False),
        'reference': result.get('reference')
    }), 200


@app.route('/pagos/paypal/order/<int:schedule_id>', methods=['POST'])
@login_required
def paypal_create_order(schedule_id):
    if session.get('role') != 'student':
        return jsonify({'ok': False, 'error': 'Solo estudiantes pueden iniciar pagos.'}), 403

    if not paypal_is_enabled():
        return jsonify({'ok': False, 'error': 'PayPal no está configurado.'}), 503

    current_user = User.query.get(session['user_id'])
    scheduled_payment = StudentPayment.query.filter_by(
        id=schedule_id,
        student_id=current_user.id,
        status='pending'
    ).first()

    if not scheduled_payment:
        return jsonify({'ok': False, 'error': 'El pago programado no está disponible.'}), 404

    access_token = get_paypal_access_token()
    if not access_token:
        return jsonify({'ok': False, 'error': 'No se pudo autenticar PayPal.'}), 502

    payload = {
        'intent': 'CAPTURE',
        'purchase_units': [{
            'reference_id': f'schedule_{scheduled_payment.id}',
            'custom_id': f'{current_user.id}:{scheduled_payment.id}',
            'description': scheduled_payment.concept,
            'amount': {
                'currency_code': scheduled_payment.currency.upper(),
                'value': f"{scheduled_payment.amount:.2f}"
            }
        }],
        'payment_source': {
            'paypal': {
                'experience_context': {
                    'payment_method_preference': 'IMMEDIATE_PAYMENT_REQUIRED',
                    'user_action': 'PAY_NOW',
                    'landing_page': 'LOGIN',
                    'shipping_preference': 'NO_SHIPPING',
                    'return_url': url_for('pagos', _external=True),
                    'cancel_url': url_for('pagos', _external=True)
                }
            }
        }
    }

    response = requests.post(
        f"{get_paypal_api_base_url()}/v2/checkout/orders",
        headers={
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {access_token}'
        },
        json=payload,
        timeout=20
    )

    if response.status_code not in [200, 201]:
        return jsonify({'ok': False, 'error': 'No se pudo crear la orden de PayPal.'}), 400

    order_data = response.json() or {}
    return jsonify({'ok': True, 'order_id': order_data.get('id')})


@app.route('/pagos/paypal/capture/<order_id>', methods=['POST'])
@login_required
def paypal_capture_order(order_id):
    if session.get('role') != 'student':
        return jsonify({'ok': False, 'error': 'Solo estudiantes pueden confirmar pagos.'}), 403

    if not paypal_is_enabled():
        return jsonify({'ok': False, 'error': 'PayPal no está configurado.'}), 503

    payload = request.get_json(silent=True) or {}
    schedule_id = payload.get('schedule_id')
    if not str(schedule_id).isdigit():
        return jsonify({'ok': False, 'error': 'Pago programado inválido.'}), 400

    current_user = User.query.get(session['user_id'])
    scheduled_payment = StudentPayment.query.filter_by(
        id=int(schedule_id),
        student_id=current_user.id,
        status='pending'
    ).first()
    if not scheduled_payment:
        return jsonify({'ok': False, 'error': 'El pago programado no está disponible.'}), 404

    access_token = get_paypal_access_token()
    if not access_token:
        return jsonify({'ok': False, 'error': 'No se pudo autenticar PayPal.'}), 502

    response = requests.post(
        f"{get_paypal_api_base_url()}/v2/checkout/orders/{order_id}/capture",
        headers={
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {access_token}'
        },
        timeout=20
    )

    if response.status_code not in [200, 201]:
        return jsonify({'ok': False, 'error': 'No se pudo capturar el pago en PayPal.'}), 400

    capture_data = response.json() or {}
    status = (capture_data.get('status') or '').upper()
    if status != 'COMPLETED':
        return jsonify({'ok': False, 'error': 'El pago no fue completado en PayPal.'}), 400

    purchase_units = capture_data.get('purchase_units') or []
    captures = (((purchase_units[0] if purchase_units else {}).get('payments') or {}).get('captures') or [])
    capture_id = captures[0].get('id') if captures else ''

    result = reconcile_paypal_capture(
        order_id=order_id,
        capture_id=capture_id,
        student_id=current_user.id,
        schedule_id=int(schedule_id)
    )

    if not result.get('ok'):
        return jsonify({'ok': False, 'error': result.get('error', 'No se pudo registrar el pago.')}), 400

    return jsonify({
        'ok': True,
        'already_processed': result.get('already_processed', False),
        'reference': result.get('reference')
    })

@app.route('/pagos/method/<int:method_id>/delete', methods=['POST'])
@login_required
def delete_payment_method(method_id):
    payment_method = PaymentMethod.query.get_or_404(method_id)

    if payment_method.user_id != session['user_id']:
        flash('No tienes permiso para eliminar este método.', 'danger')
        return redirect(url_for('pagos'))

    db.session.delete(payment_method)
    db.session.commit()
    flash('Método de pago eliminado.', 'success')
    return redirect(url_for('pagos'))

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        user = User.query.filter_by(email=email).first()
        if user:
            token = secrets.token_urlsafe(16)
            reset_tokens[token] = user.id
            link = url_for('reset_password', token=token, _external=True)
            # en un sistema real se enviaría por correo electrónico
            print(f"[DEBUG] enlace de reset para {email}: {link}")
            flash('Se ha enviado un enlace de recuperación a tu email (simulado).', 'success')
        else:
            flash('Email no encontrado.', 'danger')
        return redirect(url_for('login'))
    return render_template('forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    user_id = reset_tokens.get(token)
    if not user_id:
        flash('Enlace inválido o expirado.', 'danger')
        return redirect(url_for('login'))
    user = User.query.get(user_id)
    if request.method == 'POST':
        newpass = request.form.get('password', '')
        if len(newpass) < 6:
            flash('La contraseña debe tener al menos 6 caracteres.', 'danger')
            return redirect(url_for('reset_password', token=token))
        user.password = generate_password_hash(newpass, method='pbkdf2:sha256')
        db.session.commit()
        reset_tokens.pop(token, None)
        flash('Contraseña actualizada. Ya puedes iniciar sesión.', 'success')
        return redirect(url_for('login'))
    return render_template('reset_password.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        if not username or not password:
            flash('Por favor, completa todos los campos.', 'danger')
            return redirect(url_for('login'))
        
        user = User.query.filter_by(username=username).first()
        if not user:
            user = User.query.filter(User.email.ilike(username)).first()

        if user and not is_authorized_access_user(user):
            login_event = LoginEvent(
                user_id=user.id,
                username_attempt=username,
                success=False,
                ip_address=get_client_ip(),
                user_agent=(request.user_agent.string or '')[:255]
            )
            db.session.add(login_event)
            db.session.commit()
            flash('Acceso restringido: esta cuenta no está autorizada.', 'danger')
            return redirect(url_for('login'))
        
        if user and check_password_hash(user.password, password):
            login_event = LoginEvent(
                user_id=user.id,
                username_attempt=username,
                success=True,
                ip_address=get_client_ip(),
                user_agent=(request.user_agent.string or '')[:255]
            )
            db.session.add(login_event)
            db.session.commit()

            session['user_id'] = user.id
            session['username'] = user.username
            session['role'] = user.role
            session['access_lock_version'] = ACCESS_LOCK_VERSION
            session.permanent = True
            flash(f'¡Bienvenido {user.nombre}!', 'success')

            return redirect(url_for('index'))
        else:
            login_event = LoginEvent(
                user_id=None,
                username_attempt=username if username else '(vacío)',
                success=False,
                ip_address=get_client_ip(),
                user_agent=(request.user_agent.string or '')[:255]
            )
            db.session.add(login_event)
            db.session.commit()
            flash('Usuario o contraseña incorrectos.', 'danger')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Sesión cerrada correctamente.', 'success')
    return redirect(url_for('login'))

@app.route('/admin/login-events')
@login_required
def admin_login_events():
    current_user = User.query.get(session['user_id'])
    if not current_user or current_user.role != 'teacher':
        flash('No tienes permisos para ver este registro.', 'danger')
        return redirect(url_for('index'))

    events = LoginEvent.query.order_by(LoginEvent.created_at.desc()).limit(100).all()
    return jsonify([
        {
            'id': event.id,
            'username_attempt': event.username_attempt,
            'success': event.success,
            'ip_address': event.ip_address,
            'user_agent': event.user_agent,
            'created_at': event.created_at.strftime('%Y-%m-%d %H:%M:%S') if event.created_at else None
        }
        for event in events
    ])

# ============== RUTAS DE PERFIL ==============

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    user = User.query.get(session['user_id'])
    
    if request.method == 'POST':
        user.nombre = request.form.get('nombre', user.nombre).strip()
        user.bio = request.form.get('bio', user.bio)
        user.nationality = request.form.get('nationality', '').strip() or None
        user.preferred_payment_method = request.form.get('preferred_payment_method', '').strip() or None
        user.preferred_bank_country = request.form.get('preferred_bank_country', '').strip() or None
        new_password = request.form.get('new_password', '').strip()
        user.notification_sound_enabled = request.form.get('notification_sound_enabled') == 'on'
        
        if new_password:
            if len(new_password) < 6:
                flash('La contraseña debe tener al menos 6 caracteres.', 'danger')
                return redirect(url_for('profile'))
            user.password = generate_password_hash(new_password, method='pbkdf2:sha256')
        
        try:
            db.session.commit()
            flash('Perfil actualizado correctamente.', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Error al actualizar el perfil.', 'danger')
    
    return render_template('profile.html', user=user)

@app.route('/students')
@teacher_only
def students_profiles():
    students = User.query.filter_by(role='student').order_by(User.nombre.asc()).all()
    return render_template('students_profiles.html', students=students)

@app.route('/important-date/create', methods=['POST'])
@teacher_only
def create_important_date():
    title = request.form.get('important_title', '').strip()
    description = request.form.get('important_description', '').strip()
    meeting_link = request.form.get('important_meeting_link', '').strip()
    emoji = request.form.get('important_emoji', '').strip()
    reminder_note = request.form.get('important_reminder_note', '').strip()
    event_date_raw = request.form.get('important_event_date', '').strip()
    target_student_raw = request.form.get('important_target_student_id', '').strip()
    target_student_id = int(target_student_raw) if target_student_raw.isdigit() else None

    if not title or not event_date_raw:
        flash('Para fecha importante, completa título y fecha.', 'danger')
        return redirect(url_for('upload_material'))

    try:
        event_date = datetime.strptime(event_date_raw, '%Y-%m-%dT%H:%M')
    except ValueError:
        flash('Formato de fecha importante inválido.', 'danger')
        return redirect(url_for('upload_material'))

    event = ImportantDate(
        teacher_id=session['user_id'],
        target_student_id=target_student_id,
        title=title,
        description=description if description else None,
        meeting_link=meeting_link if meeting_link else None,
        emoji=emoji if emoji else None,
        reminder_note=reminder_note if reminder_note else None,
        event_date=event_date
    )
    db.session.add(event)
    db.session.commit()
    flash('Fecha importante creada exitosamente.', 'success')
    return redirect(url_for('upload_material'))

@app.route('/student-calendar/add', methods=['POST'])
@login_required
def add_student_calendar_note():
    if session.get('role') != 'student':
        flash('Solo estudiantes pueden agregar notas personales.', 'danger')
        return redirect(url_for('index'))

    title = request.form.get('student_note_title', '').strip()
    description = request.form.get('student_note_description', '').strip()
    emoji = request.form.get('student_note_emoji', '').strip()
    reminder_note = request.form.get('student_note_reminder', '').strip()
    event_date_raw = request.form.get('student_note_event_date', '').strip()

    if not title or not event_date_raw:
        flash('Completa título y fecha para tu nota.', 'danger')
        return redirect(url_for('classes'))

    try:
        event_date = datetime.strptime(event_date_raw, '%Y-%m-%dT%H:%M')
    except ValueError:
        flash('Formato de fecha inválido para nota personal.', 'danger')
        return redirect(url_for('classes'))

    note = StudentCalendarNote(
        student_id=session['user_id'],
        title=title,
        description=description if description else None,
        emoji=emoji if emoji else None,
        reminder_note=reminder_note if reminder_note else None,
        event_date=event_date
    )
    db.session.add(note)
    db.session.commit()
    flash('Nota personal agregada al calendario.', 'success')
    return redirect(url_for('classes'))

@app.route('/classes')
@login_required
def classes():
    if session.get('role') == 'teacher':
        return redirect(url_for('dashboard_teacher'))

    user = User.query.get(session['user_id'])

    assignments = Assignment.query.filter(
        (Assignment.target_student_id.is_(None)) | (Assignment.target_student_id == user.id)
    ).order_by(Assignment.fecha_entrega.asc()).all()

    submissions = {sub.assignment_id: sub for sub in Submission.query.filter_by(student_id=user.id).all()}

    tareas_pendientes = []
    tareas_completadas = []
    for assignment in assignments:
        if assignment.id in submissions:
            tareas_completadas.append({'tarea': assignment, 'entrega': submissions[assignment.id]})
        else:
            tareas_pendientes.append({'tarea': assignment, 'estado': 'pending' if assignment.fecha_entrega > datetime.utcnow() else 'overdue'})

    materials = Material.query.filter(
        (Material.target_student_id.is_(None)) | (Material.target_student_id == user.id)
    ).order_by(Material.fecha_creacion.desc()).all()

    important_dates = ImportantDate.query.filter(
        (ImportantDate.target_student_id.is_(None)) | (ImportantDate.target_student_id == user.id)
    ).order_by(ImportantDate.event_date.asc()).all()

    personal_notes = StudentCalendarNote.query.filter_by(student_id=user.id).order_by(StudentCalendarNote.event_date.asc()).all()

    calendar_events = []
    for item in important_dates:
        calendar_events.append({
            'title': item.title,
            'date': item.event_date.strftime('%Y-%m-%d'),
            'time': item.event_date.strftime('%H:%M'),
            'type': 'meeting' if item.meeting_link else 'important',
            'meeting_link': item.meeting_link,
            'emoji': item.emoji,
            'reminder_note': item.reminder_note
        })

    for item in materials:
        if item.scheduled_for:
            calendar_events.append({
                'title': item.titulo,
                'date': item.scheduled_for.strftime('%Y-%m-%d'),
                'time': item.scheduled_for.strftime('%H:%M'),
                'type': 'material'
            })

    for item in assignments:
        calendar_events.append({
            'title': item.titulo,
            'date': item.fecha_entrega.strftime('%Y-%m-%d'),
            'time': item.fecha_entrega.strftime('%H:%M'),
            'type': 'assignment'
        })

    for item in personal_notes:
        calendar_events.append({
            'title': item.title,
            'date': item.event_date.strftime('%Y-%m-%d'),
            'time': item.event_date.strftime('%H:%M'),
            'type': 'personal',
            'emoji': item.emoji,
            'reminder_note': item.reminder_note
        })

    return render_template(
        'classes.html',
        user=user,
        tareas_pendientes=tareas_pendientes,
        tareas_completadas=tareas_completadas,
        materials=materials,
        important_dates=important_dates,
        personal_notes=personal_notes,
        calendar_events=calendar_events
    )

@app.route('/courses', methods=['GET', 'POST'])
@login_required
def courses():
    current_user = User.query.get(session['user_id'])

    if request.method == 'POST':
        if current_user.role != 'teacher':
            flash('Solo los docentes pueden crear cursos.', 'danger')
            return redirect(url_for('courses'))

        action = request.form.get('action', 'create_course').strip()

        if action == 'create_link':
            selected_course_raw = request.form.get('course_id', '').strip()
            link_name = request.form.get('link_name', '').strip()
            link_url = request.form.get('link_url', '').strip()
            link_description = request.form.get('link_description', '').strip()

            if not selected_course_raw.isdigit():
                flash('Selecciona un curso para agregar el link.', 'danger')
                return redirect(url_for('courses'))

            selected_course = Course.query.get_or_404(int(selected_course_raw))
            if selected_course.teacher_id != current_user.id:
                flash('No tienes permiso para agregar links en ese curso.', 'danger')
                return redirect(url_for('courses'))

            if not link_name or not link_url:
                flash('Completa el nombre y la URL del link.', 'danger')
                return redirect(url_for('courses'))

            normalized_url = link_url if link_url.lower().startswith(('http://', 'https://')) else f'https://{link_url}'

            new_link = CourseLink(
                course_id=selected_course.id,
                name=link_name,
                url=normalized_url,
                description=link_description or None
            )
            db.session.add(new_link)
            db.session.commit()
            flash('Link creado correctamente.', 'success')
            return redirect(url_for('courses'))

        title = request.form.get('title', '').strip()
        description = request.form.get('description', '').strip()
        page_title = request.form.get('page_title', '').strip() or 'Hoja en blanco'
        target_student_raw = request.form.get('target_student_id', '').strip()
        target_student_id = int(target_student_raw) if target_student_raw.isdigit() else None

        if not title:
            title = f"Hoja en blanco {datetime.utcnow().strftime('%d/%m %H:%M')}"

        new_course = Course(
            title=title,
            description=description if description else None,
            teacher_id=current_user.id
        )
        db.session.add(new_course)
        db.session.flush()

        default_page_title = page_title
        default_template = {
            "time": int(datetime.utcnow().timestamp() * 1000),
            "version": "2.28.2",
            "blocks": []
        }
        initial_page = CoursePage(
            course_id=new_course.id,
            title=default_page_title,
            target_student_id=target_student_id,
            is_published=False,
            content_json=json.dumps(default_template, ensure_ascii=False)
        )
        db.session.add(initial_page)
        db.session.commit()
        flash('Borrador creado. Edita la hoja y publícala cuando esté lista.', 'success')
        return redirect(url_for('course_page_view', page_id=initial_page.id))

    if current_user.role == 'teacher':
        course_list = Course.query.filter_by(teacher_id=current_user.id).order_by(Course.created_at.desc()).all()
    else:
        course_list = Course.query.order_by(Course.created_at.desc()).all()

    teacher_ids = {course.teacher_id for course in course_list}
    teachers = User.query.filter(User.id.in_(teacher_ids)).all() if teacher_ids else []
    teacher_map = {teacher.id: teacher for teacher in teachers}

    students = User.query.filter_by(role='student').order_by(User.nombre.asc()).all() if current_user.role == 'teacher' else []

    return render_template(
        'courses.html',
        courses=course_list,
        is_teacher=current_user.role == 'teacher',
        teacher_map=teacher_map,
        students=students
    )

@app.route('/courses/<int:course_id>')
@login_required
def course_detail(course_id):
    current_user = User.query.get(session['user_id'])
    course = Course.query.get_or_404(course_id)

    if current_user.role == 'teacher' and course.teacher_id != current_user.id:
        flash('No tienes permiso para ver este curso.', 'danger')
        return redirect(url_for('courses'))

    pages_query = CoursePage.query.filter_by(course_id=course.id)
    if current_user.role == 'student':
        pages_query = pages_query.filter(
            (CoursePage.target_student_id.is_(None)) | (CoursePage.target_student_id == current_user.id)
        ).filter(CoursePage.is_published.is_(True))
    pages = pages_query.order_by(CoursePage.created_at.desc()).all()
    events = CourseCalendarEvent.query.filter_by(course_id=course.id).order_by(CourseCalendarEvent.event_date.asc()).all()
    teacher = User.query.get(course.teacher_id)
    students = User.query.filter_by(role='student').order_by(User.nombre.asc()).all() if current_user.role == 'teacher' else []
    students_map = {student.id: student for student in students}

    if current_user.role == 'teacher':
        return render_template('course_detail_teacher.html', course=course, pages=pages, events=events, teacher=teacher, students=students, students_map=students_map)

    return render_template('course_detail_student.html', course=course, pages=pages, events=events, teacher=teacher)

@app.route('/courses/<int:course_id>/links')
@login_required
def course_links(course_id):
    current_user = User.query.get(session['user_id'])
    course = Course.query.get_or_404(course_id)

    if current_user.role == 'teacher' and course.teacher_id != current_user.id:
        flash('No tienes permiso para ver este curso.', 'danger')
        return redirect(url_for('courses'))

    links = CourseLink.query.filter_by(course_id=course.id).order_by(CourseLink.created_at.desc()).all()
    is_teacher = current_user.role == 'teacher' and course.teacher_id == current_user.id

    return render_template('course_links.html', course=course, links=links, is_teacher=is_teacher)

@app.route('/courses/<int:course_id>/pages/new', methods=['POST'])
@teacher_only
def course_page_new(course_id):
    course = Course.query.get_or_404(course_id)
    if course.teacher_id != session['user_id']:
        flash('No puedes editar este curso.', 'danger')
        return redirect(url_for('courses'))

    title = request.form.get('title', '').strip() or 'Nueva clase'
    target_student_raw = request.form.get('target_student_id', '').strip()
    target_student_id = int(target_student_raw) if target_student_raw.isdigit() else None
    template_content = {
        "time": int(datetime.utcnow().timestamp() * 1000),
        "version": "2.28.2",
        "blocks": []
    }
    page = CoursePage(
        course_id=course.id,
        title=title,
        target_student_id=target_student_id,
        is_published=False,
        content_json=json.dumps(template_content, ensure_ascii=False)
    )
    db.session.add(page)
    db.session.commit()
    flash('Borrador creado. Completa el contenido y publícalo cuando quieras.', 'success')

    return redirect(url_for('course_page_view', page_id=page.id))

@app.route('/courses/<int:course_id>/events/new', methods=['POST'])
@teacher_only
def course_event_new(course_id):
    course = Course.query.get_or_404(course_id)
    if course.teacher_id != session['user_id']:
        flash('No puedes editar este curso.', 'danger')
        return redirect(url_for('courses'))

    title = request.form.get('event_title', '').strip()
    description = request.form.get('event_description', '').strip()
    event_date_raw = request.form.get('event_date', '').strip()

    if not title or not event_date_raw:
        flash('Completa título y fecha del evento.', 'danger')
        return redirect(url_for('course_detail', course_id=course.id))

    try:
        event_date = datetime.strptime(event_date_raw, '%Y-%m-%dT%H:%M')
    except ValueError:
        flash('Fecha del evento inválida.', 'danger')
        return redirect(url_for('course_detail', course_id=course.id))

    event = CourseCalendarEvent(
        course_id=course.id,
        title=title,
        description=description if description else None,
        event_date=event_date
    )
    db.session.add(event)
    db.session.commit()
    flash('Evento agregado al curso.', 'success')
    return redirect(url_for('course_detail', course_id=course.id))

@app.route('/courses/<int:course_id>/links/new', methods=['POST'])
@teacher_only
def course_link_new(course_id):
    course = Course.query.get_or_404(course_id)
    if course.teacher_id != session['user_id']:
        flash('No puedes editar este curso.', 'danger')
        return redirect(url_for('courses'))

    link_name = request.form.get('link_name', '').strip()
    link_url = request.form.get('link_url', '').strip()
    link_description = request.form.get('link_description', '').strip()

    if not link_name or not link_url:
        flash('Completa nombre y URL del enlace.', 'danger')
        return redirect(url_for('course_detail', course_id=course.id))

    # Basic URL validation
    if not (link_url.startswith('http://') or link_url.startswith('https://')):
        link_url = 'https://' + link_url

    link = CourseLink(
        course_id=course.id,
        name=link_name,
        url=link_url,
        description=link_description if link_description else None
    )
    db.session.add(link)
    db.session.commit()
    flash('Enlace agregado al curso.', 'success')
    return redirect(url_for('course_detail', course_id=course.id))

@app.route('/course-links/<int:link_id>/delete', methods=['POST'])
@teacher_only
def course_link_delete(link_id):
    link = CourseLink.query.get_or_404(link_id)
    course = Course.query.get_or_404(link.course_id)
    
    if course.teacher_id != session['user_id']:
        flash('No puedes editar este curso.', 'danger')
        return redirect(url_for('courses'))
    
    db.session.delete(link)
    db.session.commit()
    flash('Enlace eliminado del curso.', 'success')
    return redirect(url_for('course_detail', course_id=course.id))

@app.route('/course-links/<int:link_id>/edit', methods=['POST'])
@teacher_only
def course_link_edit(link_id):
    link = CourseLink.query.get_or_404(link_id)
    course = Course.query.get_or_404(link.course_id)
    
    if course.teacher_id != session['user_id']:
        flash('No puedes editar este curso.', 'danger')
        return redirect(url_for('courses'))

    link_name = request.form.get('link_name', '').strip()
    link_url = request.form.get('link_url', '').strip()
    link_description = request.form.get('link_description', '').strip()

    if not link_name or not link_url:
        flash('Completa nombre y URL del enlace.', 'danger')
        return redirect(url_for('course_detail', course_id=course.id))

    # Basic URL validation
    if not (link_url.startswith('http://') or link_url.startswith('https://')):
        link_url = 'https://' + link_url

    link.name = link_name
    link.url = link_url
    link.description = link_description if link_description else None
    db.session.commit()
    flash('Enlace actualizado.', 'success')
    return redirect(url_for('course_detail', course_id=course.id))


@app.route('/course-pages/<int:page_id>')
@login_required
def course_page_view(page_id):
    current_user = User.query.get(session['user_id'])
    page = CoursePage.query.get_or_404(page_id)
    course = Course.query.get_or_404(page.course_id)
    teacher = User.query.get(course.teacher_id)

    if current_user.role == 'teacher' and course.teacher_id != current_user.id:
        flash('No tienes permiso para ver esta clase.', 'danger')
        return redirect(url_for('courses'))

    if current_user.role == 'student' and not page.is_published:
        flash('Esta hoja aún no está publicada.', 'warning')
        return redirect(url_for('course_detail', course_id=course.id))

    if current_user.role == 'student' and page.target_student_id and page.target_student_id != current_user.id:
        flash('Esta hoja no está asignada para ti.', 'danger')
        return redirect(url_for('course_detail', course_id=course.id))

    note = None
    if current_user.role == 'student':
        note = CourseStudentNote.query.filter_by(page_id=page.id, student_id=current_user.id).first()
        if not note:
            note = CourseStudentNote(page_id=page.id, student_id=current_user.id)
            db.session.add(note)
            db.session.commit()

    questions = CourseQuestion.query.filter_by(page_id=page.id).order_by(CourseQuestion.created_at.desc()).all()
    question_ids = [question.id for question in questions]
    answers = CourseAnswer.query.filter(CourseAnswer.question_id.in_(question_ids)).all() if question_ids else []

    answers_by_q = {}
    for answer in answers:
        answers_by_q.setdefault(answer.question_id, []).append(answer)

    student_ids = {question.student_id for question in questions}
    teacher_ids = {answer.teacher_id for answer in answers}

    student_users = User.query.filter(User.id.in_(student_ids)).all() if student_ids else []
    teacher_users = User.query.filter(User.id.in_(teacher_ids)).all() if teacher_ids else []

    student_map = {user.id: user for user in student_users}
    teacher_map = {user.id: user for user in teacher_users}

    return render_template(
        'course_page_workspace.html',
        page=page,
        course=course,
        teacher=teacher,
        note=note,
        questions=questions,
        answers_by_q=answers_by_q,
        student_map=student_map,
        teacher_map=teacher_map,
        is_teacher=current_user.role == 'teacher'
    )

@app.route('/course-pages/<int:page_id>/save', methods=['POST'])
@teacher_only
def course_page_save_teacher(page_id):
    page = CoursePage.query.get_or_404(page_id)
    course = Course.query.get_or_404(page.course_id)

    if course.teacher_id != session['user_id']:
        return jsonify({'ok': False, 'error': 'No autorizado'}), 403

    data = request.get_json(silent=True)
    if not data:
        return jsonify({'ok': False, 'error': 'JSON inválido'}), 400

    page.content_json = json.dumps(data, ensure_ascii=False)
    page.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/course-pages/<int:page_id>/publish', methods=['POST'])
@teacher_only
def course_page_publish(page_id):
    page = CoursePage.query.get_or_404(page_id)
    course = Course.query.get_or_404(page.course_id)

    if course.teacher_id != session['user_id']:
        return jsonify({'ok': False, 'error': 'No autorizado'}), 403

    was_published = bool(page.is_published)
    page.is_published = True
    page.updated_at = datetime.utcnow()

    if not was_published:
        teacher = User.query.get(session['user_id'])
        teacher_name = teacher.nombre or teacher.username
        notify_course_page_students(page, teacher_name)

    db.session.commit()

    wants_json = request.is_json or ('application/json' in (request.headers.get('Accept') or ''))
    if wants_json:
        return jsonify({'ok': True})

    next_url = request.form.get('next_url', '').strip()
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect(url_for('course_page_view', page_id=page.id))

@app.route('/course-pages/<int:page_id>/unpublish', methods=['POST'])
@teacher_only
def course_page_unpublish(page_id):
    page = CoursePage.query.get_or_404(page_id)
    course = Course.query.get_or_404(page.course_id)

    if course.teacher_id != session['user_id']:
        return jsonify({'ok': False, 'error': 'No autorizado'}), 403

    page.is_published = False
    page.updated_at = datetime.utcnow()
    db.session.commit()

    wants_json = request.is_json or ('application/json' in (request.headers.get('Accept') or ''))
    if wants_json:
        return jsonify({'ok': True})

    next_url = request.form.get('next_url', '').strip()
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect(url_for('course_page_view', page_id=page.id))

@app.route('/course-pages/<int:page_id>/delete', methods=['POST'])
@teacher_only
def course_page_delete(page_id):
    page = CoursePage.query.get_or_404(page_id)
    course = Course.query.get_or_404(page.course_id)

    if course.teacher_id != session['user_id']:
        flash('No autorizado para eliminar esta hoja.', 'danger')
        return redirect(url_for('courses'))

    question_ids = [q.id for q in CourseQuestion.query.filter_by(page_id=page.id).all()]
    if question_ids:
        CourseAnswer.query.filter(CourseAnswer.question_id.in_(question_ids)).delete(synchronize_session=False)

    CourseQuestion.query.filter_by(page_id=page.id).delete(synchronize_session=False)
    CourseStudentNote.query.filter_by(page_id=page.id).delete(synchronize_session=False)
    db.session.delete(page)
    db.session.commit()

    flash('Hoja eliminada correctamente.', 'success')
    return redirect(url_for('course_detail', course_id=course.id))

@app.route('/courses/<int:course_id>/delete', methods=['POST'])
@teacher_only
def course_delete(course_id):
    course = Course.query.get_or_404(course_id)

    if course.teacher_id != session['user_id']:
        flash('No autorizado para eliminar este curso.', 'danger')
        return redirect(url_for('courses'))

    page_ids = [page.id for page in CoursePage.query.filter_by(course_id=course.id).all()]

    if page_ids:
        question_ids = [q.id for q in CourseQuestion.query.filter(CourseQuestion.page_id.in_(page_ids)).all()]
        if question_ids:
            CourseAnswer.query.filter(CourseAnswer.question_id.in_(question_ids)).delete(synchronize_session=False)

        CourseQuestion.query.filter(CourseQuestion.page_id.in_(page_ids)).delete(synchronize_session=False)
        CourseStudentNote.query.filter(CourseStudentNote.page_id.in_(page_ids)).delete(synchronize_session=False)
        CoursePage.query.filter(CoursePage.id.in_(page_ids)).delete(synchronize_session=False)

    CourseCalendarEvent.query.filter_by(course_id=course.id).delete(synchronize_session=False)
    db.session.delete(course)
    db.session.commit()

    flash('Curso eliminado correctamente.', 'success')
    return redirect(url_for('courses'))

@app.route('/course-pages/<int:page_id>/images/upload', methods=['POST'])
@teacher_only
def course_page_upload_image(page_id):
    page = CoursePage.query.get_or_404(page_id)
    course = Course.query.get_or_404(page.course_id)

    if course.teacher_id != session['user_id']:
        return jsonify({'ok': False, 'error': 'No autorizado'}), 403

    if 'image' not in request.files:
        return jsonify({'ok': False, 'error': 'No se recibió archivo'}), 400

    image_file = request.files['image']
    if not image_file or image_file.filename == '':
        return jsonify({'ok': False, 'error': 'Archivo vacío'}), 400

    if not allowed_file(image_file.filename):
        return jsonify({'ok': False, 'error': 'Tipo de archivo no permitido'}), 400

    extension = image_file.filename.rsplit('.', 1)[1].lower()
    if extension not in {'png', 'jpg', 'jpeg', 'gif'}:
        return jsonify({'ok': False, 'error': 'Solo imágenes (png, jpg, jpeg, gif)'}), 400

    safe_name = secure_filename(image_file.filename)
    filename = secure_filename(f"course_page_{page.id}_{int(datetime.utcnow().timestamp())}_{safe_name}")
    storage_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    image_file.save(storage_path)

    file_url = url_for('download_uploaded_file', filename=filename)
    return jsonify({'ok': True, 'url': file_url})

@app.route('/uploads/<path:filename>')
@login_required
def download_uploaded_file(filename):
    safe_name = secure_filename(filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_name)

    if not os.path.exists(file_path):
        return jsonify({'ok': False, 'error': 'Archivo no encontrado'}), 404

    return send_file(file_path, as_attachment=False)

@app.route('/course-pages/<int:page_id>/notes/save', methods=['POST'])
@login_required
def course_save_notes(page_id):
    if session.get('role') != 'student':
        return jsonify({'ok': False, 'error': 'Solo estudiantes'}), 403

    note = CourseStudentNote.query.filter_by(page_id=page_id, student_id=session['user_id']).first()
    if not note:
        return jsonify({'ok': False, 'error': 'Nota no encontrada'}), 404

    data = request.get_json(silent=True)
    if not data:
        return jsonify({'ok': False, 'error': 'JSON inválido'}), 400

    note.notes_json = json.dumps(data, ensure_ascii=False)
    note.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/course-pages/<int:page_id>/questions/new', methods=['POST'])
@login_required
def course_new_question(page_id):
    if session.get('role') != 'student':
        flash('Solo estudiantes pueden enviar preguntas.', 'danger')
        return redirect(url_for('course_page_view', page_id=page_id))

    text_value = request.form.get('text', '').strip()
    if not text_value:
        flash('Escribe tu pregunta o comentario.', 'danger')
        return redirect(url_for('course_page_view', page_id=page_id))

    question = CourseQuestion(page_id=page_id, student_id=session['user_id'], text=text_value)
    db.session.add(question)
    db.session.commit()
    flash('Pregunta enviada al docente.', 'success')
    return redirect(url_for('course_page_view', page_id=page_id))

@app.route('/course-questions/<int:question_id>/answer', methods=['POST'])
@teacher_only
def course_answer_question(question_id):
    question = CourseQuestion.query.get_or_404(question_id)
    page = CoursePage.query.get_or_404(question.page_id)
    course = Course.query.get_or_404(page.course_id)

    if course.teacher_id != session['user_id']:
        flash('No puedes responder en este curso.', 'danger')
        return redirect(url_for('courses'))

    text_value = request.form.get('text', '').strip()
    if not text_value:
        flash('Escribe una respuesta.', 'danger')
        return redirect(url_for('course_page_view', page_id=page.id))

    answer = CourseAnswer(question_id=question.id, teacher_id=session['user_id'], text=text_value)
    db.session.add(answer)
    db.session.commit()
    flash('Respuesta enviada.', 'success')
    return redirect(url_for('course_page_view', page_id=page.id))

# ============== RUTAS DEL DASHBOARD ==============

@app.route('/dashboard/teacher')
@teacher_only
def dashboard_teacher():
    user = User.query.get(session['user_id'])
    tareas = Assignment.query.filter_by(teacher_id=user.id).order_by(Assignment.fecha_creacion.desc()).all()
    materiales = Material.query.filter_by(user_id=user.id).order_by(Material.fecha_creacion.desc()).all()
    daily_verse = get_daily_verse(session.get('lang', DEFAULT_LANGUAGE))
    
    total_tareas = len(tareas)
    total_estudiantes = len(set([sub.student_id for tarea in tareas for sub in tarea.entregas]))
    total_entregas = sum(len(tarea.entregas) for tarea in tareas)
    
    return render_template('dashboard_teacher.html', 
                         user=user, 
                         tareas=tareas, 
                         materiales=materiales,
                         total_tareas=total_tareas,
                         total_estudiantes=total_estudiantes,
                         total_entregas=total_entregas,
                         daily_verse=daily_verse)

@app.route('/dashboard/student')
@login_required
def dashboard_student():
    if session.get('role') == 'teacher':
        return redirect(url_for('dashboard_teacher'))
    
    user = User.query.get(session['user_id'])
    
    # Obtener tareas para todos o específicas del estudiante
    todas_tareas = Assignment.query.filter(
        (Assignment.target_student_id.is_(None)) | (Assignment.target_student_id == user.id)
    ).all()
    
    # Obtener entregas del estudiante
    entregas_dict = {sub.assignment_id: sub for sub in Submission.query.filter_by(student_id=user.id).all()}
    
    # Organizar tareas
    tareas_pendientes = []
    tareas_completadas = []
    
    for tarea in todas_tareas:
        if tarea.id in entregas_dict:
            tareas_completadas.append({
                'tarea': tarea,
                'entrega': entregas_dict[tarea.id],
                'estado': 'completada'
            })
        else:
            tareas_pendientes.append({
                'tarea': tarea,
                'entrega': None,
                'estado': 'pendiente' if tarea.fecha_entrega > datetime.utcnow() else 'vencida'
            })
    
    materiales = Material.query.filter(
        (Material.target_student_id.is_(None)) | (Material.target_student_id == user.id)
    ).order_by(Material.fecha_creacion.desc()).all()
    daily_verse = get_daily_verse(session.get('lang', DEFAULT_LANGUAGE))
    show_welcome_popup = request.args.get('welcome') == '1'
    
    return render_template('dashboard_student.html',
                         user=user,
                         tareas_pendientes=tareas_pendientes,
                         tareas_completadas=tareas_completadas,
                         materiales=materiales,
                         daily_verse=daily_verse,
                         show_welcome_popup=show_welcome_popup)

# ============== MENSAJES Y NOTIFICACIONES ==============

@app.route('/messages', methods=['GET', 'POST'])
@login_required
def messages():
    current_user = User.query.get(session['user_id'])

    if request.method == 'POST':
        recipient_id = request.form.get('recipient_id', '').strip()
        content = request.form.get('content', '').strip()

        if not recipient_id.isdigit() or not content:
            flash('Selecciona destinatario y escribe un mensaje.', 'danger')
            return redirect(url_for('messages'))

        recipient = User.query.get(int(recipient_id))
        if not recipient:
            flash('Destinatario no válido.', 'danger')
            return redirect(url_for('messages'))

        if recipient.id == current_user.id:
            flash('No puedes enviarte mensajes a ti mismo.', 'warning')
            return redirect(url_for('messages'))

        if recipient.role == current_user.role:
            flash('Los mensajes deben ser entre docente y estudiante.', 'warning')
            return redirect(url_for('messages'))

        new_message = Message(
            sender_id=current_user.id,
            recipient_id=recipient.id,
            content=content
        )
        db.session.add(new_message)

        new_notification = UserNotification(
            user_id=recipient.id,
            title='Nuevo mensaje',
            body=f'{current_user.nombre or current_user.username}: {content[:90]}',
            link=url_for('messages')
        )
        db.session.add(new_notification)
        db.session.commit()

        flash('Mensaje enviado correctamente.', 'success')
        return redirect(url_for('messages'))

    received_messages = Message.query.filter_by(recipient_id=current_user.id) \
        .order_by(Message.created_at.desc()).limit(100).all()
    sent_messages = Message.query.filter_by(sender_id=current_user.id) \
        .order_by(Message.created_at.desc()).limit(100).all()

    Message.query.filter_by(recipient_id=current_user.id, is_read=False).update({'is_read': True})
    UserNotification.query.filter_by(user_id=current_user.id, is_read=False).update({'is_read': True})
    db.session.commit()

    recipients = User.query.filter(User.role != current_user.role).order_by(User.nombre.asc(), User.username.asc()).all()
    users_map = {user.id: user for user in User.query.filter(User.id.in_({msg.sender_id for msg in received_messages} | {msg.recipient_id for msg in sent_messages})).all()} if (received_messages or sent_messages) else {}

    return render_template(
        'messages.html',
        user=current_user,
        recipients=recipients,
        received_messages=received_messages,
        sent_messages=sent_messages,
        users_map=users_map
    )

@app.route('/notifications/mark-read', methods=['POST'])
@login_required
def mark_notifications_read():
    current_user_id = session['user_id']
    Message.query.filter_by(recipient_id=current_user_id, is_read=False).update({'is_read': True})
    UserNotification.query.filter_by(user_id=current_user_id, is_read=False).update({'is_read': True})
    db.session.commit()
    return jsonify({'ok': True})

@app.route('/notifications/sound-preference', methods=['POST'])
@login_required
def update_notification_sound_preference():
    user = User.query.get(session['user_id'])
    data = request.get_json(silent=True) or {}
    enabled = bool(data.get('enabled', True))
    user.notification_sound_enabled = enabled
    db.session.commit()
    return jsonify({'ok': True, 'notification_sound_enabled': user.notification_sound_enabled})

# ============== RUTAS DE MATERIALES ==============

@app.route('/material/upload', methods=['GET', 'POST'])
@teacher_only
def upload_material():
    if request.method == 'POST':
        titulo = request.form.get('titulo', '').strip()
        descripcion = request.form.get('descripcion', '').strip()
        tipo = request.form.get('tipo', '')
        board_status = request.form.get('board_status', 'published').strip().lower()
        scheduled_for_raw = request.form.get('scheduled_for', '').strip()
        planning_items = request.form.get('planning_items', '').strip()
        target_student_raw = request.form.get('target_student_id', '').strip()
        target_student_id = int(target_student_raw) if target_student_raw.isdigit() else None

        if board_status not in ['draft', 'planned', 'published']:
            board_status = 'published'

        scheduled_for = None
        if scheduled_for_raw:
            try:
                scheduled_for = datetime.strptime(scheduled_for_raw, '%Y-%m-%dT%H:%M')
            except ValueError:
                flash('Formato de fecha de calendario incorrecto.', 'danger')
                return redirect(url_for('upload_material'))
        
        if not titulo or not tipo:
            flash('Por favor, completa todos los campos.', 'danger')
            return redirect(url_for('upload_material'))
        
        try:
            if tipo == 'video':
                url_video = request.form.get('url_video', '').strip()
                if not url_video:
                    flash('Por favor, proporciona una URL de video.', 'danger')
                    return redirect(url_for('upload_material'))
                
                material = Material(
                    titulo=titulo,
                    descripcion=descripcion,
                    tipo=tipo,
                    board_status=board_status,
                    scheduled_for=scheduled_for,
                    planning_items=planning_items if planning_items else None,
                    url_video=url_video,
                    user_id=session['user_id'],
                    target_student_id=target_student_id
                )
            else:
                if 'archivo' not in request.files:
                    flash('Por favor, selecciona un archivo.', 'danger')
                    return redirect(url_for('upload_material'))
                
                archivo = request.files['archivo']
                if archivo.filename == '':
                    flash('Por favor, selecciona un archivo.', 'danger')
                    return redirect(url_for('upload_material'))
                
                if not allowed_file(archivo.filename):
                    flash('Tipo de archivo no permitido.', 'danger')
                    return redirect(url_for('upload_material'))
                
                filename = secure_filename(f"{datetime.utcnow().timestamp()}_{archivo.filename}")
                archivo.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                word_content = None
                if tipo == 'documento' and filename.lower().endswith('.docx'):
                    word_content = extract_docx_text(file_path)
                
                material = Material(
                    titulo=titulo,
                    descripcion=descripcion,
                    tipo=tipo,
                    board_status=board_status,
                    scheduled_for=scheduled_for,
                    planning_items=planning_items if planning_items else None,
                    word_content=word_content,
                    url_archivo=filename,
                    user_id=session['user_id'],
                    target_student_id=target_student_id
                )
            
            db.session.add(material)
            db.session.commit()
            flash('Material subido exitosamente.', 'success')
            return redirect(url_for('dashboard_teacher'))
        
        except Exception as e:
            db.session.rollback()
            flash('Error al subir el material.', 'danger')
    
    students = User.query.filter_by(role='student').order_by(User.nombre.asc()).all()
    scheduled_materials = Material.query.filter_by(user_id=session['user_id']) \
        .filter(Material.scheduled_for.isnot(None)) \
        .order_by(Material.scheduled_for.asc()) \
        .limit(30).all()
    teacher_materials = Material.query.filter_by(user_id=session['user_id']) \
        .order_by(Material.fecha_creacion.desc()) \
        .limit(20).all()
    teacher_assignments = Assignment.query.filter_by(teacher_id=session['user_id']) \
        .order_by(Assignment.fecha_creacion.desc()) \
        .limit(20).all()
    important_dates = ImportantDate.query.filter_by(teacher_id=session['user_id']).order_by(ImportantDate.event_date.asc()).limit(60).all()
    return render_template(
        'upload_material.html',
        students=students,
        scheduled_materials=scheduled_materials,
        important_dates=important_dates,
        teacher_materials=teacher_materials,
        teacher_assignments=teacher_assignments
    )

@app.route('/material/<int:material_id>/download')
@login_required
def download_material(material_id):
    material = Material.query.get_or_404(material_id)

    if session.get('role') == 'student' and material.target_student_id and material.target_student_id != session['user_id']:
        flash('No tienes permiso para descargar este material.', 'danger')
        return redirect(url_for('dashboard_student'))
    
    if material.url_archivo:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], material.url_archivo)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
    
    flash('Archivo no encontrado.', 'danger')
    return redirect(url_for('dashboard_student'))

@app.route('/material/<int:material_id>/view')
@login_required
def view_material(material_id):
    material = Material.query.get_or_404(material_id)

    if session.get('role') == 'student' and material.target_student_id and material.target_student_id != session['user_id']:
        flash('No tienes permiso para ver este material.', 'danger')
        return redirect(url_for('classes'))

    if material.url_archivo:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], material.url_archivo)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=False)

    flash('Archivo no encontrado.', 'danger')
    return redirect(url_for('classes'))

@app.route('/material/<int:material_id>/workspace')
@login_required
def material_workspace(material_id):
    material = Material.query.get_or_404(material_id)
    user = User.query.get(session['user_id'])

    if user.role == 'teacher' and material.user_id != user.id:
        flash('No tienes permiso para ver este material.', 'danger')
        return redirect(url_for('dashboard_teacher'))

    if user.role == 'student' and material.target_student_id and material.target_student_id != user.id:
        flash('No tienes permiso para ver este material.', 'danger')
        return redirect(url_for('classes'))

    file_path = None
    if material.url_archivo:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], material.url_archivo)
    is_pdf_file = bool(material.url_archivo and material.url_archivo.lower().endswith('.pdf'))

    text_content = material.word_content or ''
    if not text_content and material.url_archivo and file_path and os.path.exists(file_path):
        lower_name = material.url_archivo.lower()
        if lower_name.endswith('.docx'):
            text_content = extract_docx_text(file_path)
            material.word_content = text_content
            db.session.commit()
        elif lower_name.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as txt_file:
                text_content = txt_file.read()

    selected_student_id = None
    student_options = []
    annotations_query = MaterialAnnotation.query.filter_by(material_id=material.id)

    if user.role == 'student':
        selected_student_id = user.id
        annotations_query = annotations_query.filter_by(student_id=user.id)
    else:
        students_query = User.query.filter_by(role='student').order_by(User.nombre.asc(), User.username.asc())
        if material.target_student_id:
            students_query = students_query.filter(User.id == material.target_student_id)
        student_options = students_query.all()

        selected_student_id = request.args.get('student_id', type=int)
        valid_ids = {student.id for student in student_options}
        if selected_student_id not in valid_ids:
            with_notes = MaterialAnnotation.query.filter_by(material_id=material.id) \
                .order_by(MaterialAnnotation.created_at.desc()).all()
            for note in with_notes:
                if note.student_id in valid_ids:
                    selected_student_id = note.student_id
                    break
            if selected_student_id not in valid_ids and student_options:
                selected_student_id = student_options[0].id

        if selected_student_id:
            annotations_query = annotations_query.filter_by(student_id=selected_student_id)
        else:
            annotations_query = annotations_query.filter(text('1=0'))

    annotations = annotations_query.order_by(MaterialAnnotation.created_at.asc()).all()
    student_map = {student.id: student for student in User.query.filter_by(role='student').all()}

    annotations_payload = [
        {
            'id': note.id,
            'student_id': note.student_id,
            'student_name': (student_map[note.student_id].nombre if note.student_id in student_map and student_map[note.student_id].nombre else (student_map[note.student_id].username if note.student_id in student_map else 'Estudiante')),
            'selected_text': note.selected_text,
            'comment': note.comment,
            'mark_type': note.mark_type,
            'start_offset': note.start_offset,
            'end_offset': note.end_offset,
            'page_number': note.page_number,
            'created_at': note.created_at.strftime('%Y-%m-%d %H:%M')
        }
        for note in annotations
    ]

    return render_template(
        'material_workspace.html',
        material=material,
        user=user,
        text_content=text_content,
        annotations=annotations_payload,
        selected_student_id=selected_student_id,
        student_options=student_options,
        file_view_url=url_for('view_material', material_id=material.id),
        can_annotate=(user.role == 'student' and (bool(text_content.strip()) or is_pdf_file)),
        is_pdf_file=is_pdf_file
    )

@app.route('/material/<int:material_id>/annotations', methods=['GET', 'POST'])
@login_required
def material_annotations_api(material_id):
    material = Material.query.get_or_404(material_id)
    user = User.query.get(session['user_id'])

    if user.role == 'teacher' and material.user_id != user.id:
        return jsonify({'error': 'No autorizado'}), 403

    if user.role == 'student' and material.target_student_id and material.target_student_id != user.id:
        return jsonify({'error': 'No autorizado'}), 403

    if request.method == 'GET':
        annotations_query = MaterialAnnotation.query.filter_by(material_id=material.id)

        if user.role == 'student':
            annotations_query = annotations_query.filter_by(student_id=user.id)
        else:
            selected_student_id = request.args.get('student_id', type=int)
            if selected_student_id:
                annotations_query = annotations_query.filter_by(student_id=selected_student_id)

        notes = annotations_query.order_by(MaterialAnnotation.created_at.asc()).all()
        student_map = {student.id: student for student in User.query.filter_by(role='student').all()}
        payload = [
            {
                'id': note.id,
                'student_id': note.student_id,
                'student_name': (student_map[note.student_id].nombre if note.student_id in student_map and student_map[note.student_id].nombre else (student_map[note.student_id].username if note.student_id in student_map else 'Estudiante')),
                'selected_text': note.selected_text,
                'comment': note.comment,
                'mark_type': note.mark_type,
                'start_offset': note.start_offset,
                'end_offset': note.end_offset,
                'page_number': note.page_number,
                'created_at': note.created_at.strftime('%Y-%m-%d %H:%M')
            }
            for note in notes
        ]
        return jsonify({'annotations': payload})

    if user.role != 'student':
        return jsonify({'error': 'Solo los estudiantes pueden agregar anotaciones'}), 403

    data = request.get_json(silent=True) or {}
    selected_text = (data.get('selected_text') or '').strip()
    comment = (data.get('comment') or '').strip()
    mark_type = (data.get('mark_type') or 'highlight').strip().lower()
    page_number = data.get('page_number')

    try:
        start_offset = int(data.get('start_offset'))
        end_offset = int(data.get('end_offset'))
    except (TypeError, ValueError):
        return jsonify({'error': 'Offsets inválidos'}), 400

    if page_number is not None:
        try:
            page_number = int(page_number)
        except (TypeError, ValueError):
            page_number = None

    if mark_type not in ['highlight', 'underline']:
        mark_type = 'highlight'

    if not selected_text or end_offset <= start_offset or start_offset < 0:
        return jsonify({'error': 'Selecciona texto válido para anotar'}), 400

    note = MaterialAnnotation(
        material_id=material.id,
        student_id=user.id,
        selected_text=selected_text,
        comment=comment if comment else None,
        mark_type=mark_type,
        start_offset=start_offset,
        end_offset=end_offset,
        page_number=page_number
    )
    db.session.add(note)
    db.session.commit()

    return jsonify({'ok': True, 'annotation_id': note.id})

@app.route('/material/<int:material_id>/edit-word', methods=['GET', 'POST'])
@teacher_only
def edit_word_material(material_id):
    material = Material.query.get_or_404(material_id)

    if material.user_id != session['user_id']:
        flash('No tienes permiso para editar este material.', 'danger')
        return redirect(url_for('dashboard_teacher'))

    if material.tipo != 'documento' or not material.url_archivo or not material.url_archivo.lower().endswith('.docx'):
        flash('Solo los archivos Word (.docx) se pueden editar desde la plataforma.', 'warning')
        return redirect(url_for('dashboard_teacher'))

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], material.url_archivo)
    if not os.path.exists(file_path):
        flash('Archivo Word no encontrado.', 'danger')
        return redirect(url_for('dashboard_teacher'))

    if request.method == 'POST':
        content = request.form.get('word_content', '').strip()
        save_docx_text(file_path, content)
        material.word_content = content
        db.session.commit()
        flash('Documento Word actualizado exitosamente.', 'success')
        return redirect(url_for('dashboard_teacher'))

    if not material.word_content:
        material.word_content = extract_docx_text(file_path)
        db.session.commit()

    return render_template('edit_word_material.html', material=material)

@app.route('/material/<int:material_id>/delete', methods=['POST'])
@teacher_only
def delete_material(material_id):
    material = Material.query.get_or_404(material_id)
    
    if material.user_id != session['user_id']:
        flash('No tienes permiso para eliminar este material.', 'danger')
        return redirect(url_for('dashboard_teacher'))
    
    if material.url_archivo:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], material.url_archivo)
        if os.path.exists(file_path):
            os.remove(file_path)
    
    try:
        db.session.delete(material)
        db.session.commit()
        flash('Material eliminado correctamente.', 'success')
    except:
        db.session.rollback()
        flash('Error al eliminar el material.', 'danger')

    next_url = request.form.get('next_url', '').strip()
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect(request.referrer or url_for('dashboard_teacher'))

# ============== RUTAS DE TAREAS ==============

@app.route('/assignment/create', methods=['GET', 'POST'])
@teacher_only
def create_assignment():
    if request.method == 'POST':
        titulo = request.form.get('titulo', '').strip()
        descripcion = request.form.get('descripcion', '').strip()
        fecha_entrega = request.form.get('fecha_entrega', '')
        target_student_raw = request.form.get('target_student_id', '').strip()
        target_student_id = int(target_student_raw) if target_student_raw.isdigit() else None
        
        if not titulo or not descripcion or not fecha_entrega:
            flash('Por favor, completa todos los campos.', 'danger')
            return redirect(url_for('create_assignment'))
        
        try:
            fecha_entrega_dt = datetime.strptime(fecha_entrega, '%Y-%m-%dT%H:%M')
            
            assignment = Assignment(
                titulo=titulo,
                descripcion=descripcion,
                fecha_entrega=fecha_entrega_dt,
                teacher_id=session['user_id'],
                target_student_id=target_student_id
            )
            
            db.session.add(assignment)
            db.session.flush()

            teacher = User.query.get(session['user_id'])
            teacher_name = teacher.nombre or teacher.username
            notify_assignment_students(assignment, teacher_name, mode='creada')

            db.session.commit()
            flash('Tarea creada exitosamente.', 'success')
            return redirect(url_for('dashboard_teacher'))
        
        except ValueError:
            flash('Formato de fecha incorrecto.', 'danger')
        except Exception as e:
            db.session.rollback()
            flash('Error al crear la tarea.', 'danger')
    
    students = User.query.filter_by(role='student').order_by(User.nombre.asc()).all()
    return render_template('create_assignment.html', students=students)

@app.route('/assignment/<int:assignment_id>')
@login_required
def assignment_detail(assignment_id):
    assignment = Assignment.query.get_or_404(assignment_id)
    
    if session.get('role') == 'teacher' and assignment.teacher_id != session['user_id']:
        flash('No tienes permiso para ver esta tarea.', 'danger')
        return redirect(url_for('dashboard_teacher'))

    if session.get('role') == 'student' and assignment.target_student_id and assignment.target_student_id != session['user_id']:
        flash('No tienes permiso para ver esta tarea.', 'danger')
        return redirect(url_for('dashboard_student'))
    
    user = User.query.get(session['user_id'])
    
    # Obtener mi entrega si soy estudiante
    mi_entrega = None
    if session.get('role') == 'student':
        mi_entrega = Submission.query.filter_by(
            assignment_id=assignment_id,
            student_id=session['user_id']
        ).first()
    
    # Obtener todas las entregas si soy maestro
    entregas = []
    if session.get('role') == 'teacher':
        entregas = Submission.query.filter_by(assignment_id=assignment_id).all()
    
    return render_template('assignment_detail.html',
                         assignment=assignment,
                         mi_entrega=mi_entrega,
                         entregas=entregas,
                         user=user)

@app.route('/assignment/<int:assignment_id>/edit', methods=['GET', 'POST'])
@teacher_only
def edit_assignment(assignment_id):
    assignment = Assignment.query.get_or_404(assignment_id)
    
    if assignment.teacher_id != session['user_id']:
        flash('No tienes permiso para editar esta tarea.', 'danger')
        return redirect(url_for('dashboard_teacher'))
    
    if request.method == 'POST':
        previous_target = assignment.target_student_id
        assignment.titulo = request.form.get('titulo', assignment.titulo).strip()
        assignment.descripcion = request.form.get('descripcion', assignment.descripcion).strip()
        target_student_raw = request.form.get('target_student_id', '').strip()
        assignment.target_student_id = int(target_student_raw) if target_student_raw.isdigit() else None
        fecha_entrega = request.form.get('fecha_entrega', '')
        
        try:
            if fecha_entrega:
                assignment.fecha_entrega = datetime.strptime(fecha_entrega, '%Y-%m-%dT%H:%M')

            if previous_target != assignment.target_student_id:
                teacher = User.query.get(session['user_id'])
                teacher_name = teacher.nombre or teacher.username
                notify_assignment_students(assignment, teacher_name, mode='actualizada')
            
            db.session.commit()
            flash('Tarea actualizada correctamente.', 'success')
            return redirect(url_for('assignment_detail', assignment_id=assignment_id))
        except ValueError:
            flash('Formato de fecha incorrecto.', 'danger')
        except:
            db.session.rollback()
            flash('Error al actualizar la tarea.', 'danger')
    
    students = User.query.filter_by(role='student').order_by(User.nombre.asc()).all()
    return render_template('create_assignment.html', assignment=assignment, edit_mode=True, students=students)

@app.route('/assignment/<int:assignment_id>/delete', methods=['POST'])
@teacher_only
def delete_assignment(assignment_id):
    assignment = Assignment.query.get_or_404(assignment_id)
    
    if assignment.teacher_id != session['user_id']:
        flash('No tienes permiso para eliminar esta tarea.', 'danger')
        return redirect(url_for('dashboard_teacher'))
    
    try:
        db.session.delete(assignment)
        db.session.commit()
        flash('Tarea eliminada correctamente.', 'success')
    except:
        db.session.rollback()
        flash('Error al eliminar la tarea.', 'danger')

    next_url = request.form.get('next_url', '').strip()
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect(request.referrer or url_for('dashboard_teacher'))

# ============== RUTAS DE ENTREGAS ==============

@app.route('/submission/<int:assignment_id>', methods=['POST'])
@login_required
def submit_assignment(assignment_id):
    if session.get('role') != 'student':
        flash('Solo los estudiantes pueden entregar tareas.', 'danger')
        return redirect(url_for('dashboard_teacher'))
    
    assignment = Assignment.query.get_or_404(assignment_id)
    
    if 'archivo' not in request.files:
        flash('Por favor, selecciona un archivo.', 'danger')
        return redirect(url_for('assignment_detail', assignment_id=assignment_id))
    
    archivo = request.files['archivo']
    
    if archivo.filename == '':
        flash('Por favor, selecciona un archivo.', 'danger')
        return redirect(url_for('assignment_detail', assignment_id=assignment_id))
    
    if not allowed_file(archivo.filename):
        flash('Tipo de archivo no permitido.', 'danger')
        return redirect(url_for('assignment_detail', assignment_id=assignment_id))
    
    try:
        # Eliminar entrega anterior si existe
        entrega_anterior = Submission.query.filter_by(
            assignment_id=assignment_id,
            student_id=session['user_id']
        ).first()
        
        if entrega_anterior and entrega_anterior.url_archivo:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], entrega_anterior.url_archivo)
            if os.path.exists(file_path):
                os.remove(file_path)
            db.session.delete(entrega_anterior)
            db.session.commit()
        
        filename = secure_filename(f"{datetime.utcnow().timestamp()}_{archivo.filename}")
        archivo.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        
        nueva_entrega = Submission(
            assignment_id=assignment_id,
            student_id=session['user_id'],
            url_archivo=filename
        )
        
        db.session.add(nueva_entrega)
        db.session.commit()
        flash('Tarea entregada exitosamente.', 'success')
    except Exception as e:
        db.session.rollback()
        flash('Error al entregar la tarea.', 'danger')
    
    return redirect(url_for('assignment_detail', assignment_id=assignment_id))

@app.route('/submission/<int:submission_id>/grade', methods=['POST'])
@teacher_only
def grade_submission(submission_id):
    submission = Submission.query.get_or_404(submission_id)
    assignment = submission.tarea
    
    if assignment.teacher_id != session['user_id']:
        flash('No tienes permiso para calificar esta entrega.', 'danger')
        return redirect(url_for('dashboard_teacher'))
    
    try:
        calificacion = int(request.form.get('calificacion', 0))
        if calificacion < 0 or calificacion > 100:
            calificacion = 0
        
        submission.calificacion = calificacion
        submission.comentario = request.form.get('comentario', '').strip()
        
        db.session.commit()
        flash('Entrega calificada correctamente.', 'success')
    except:
        db.session.rollback()
        flash('Error al calificar la entrega.', 'danger')
    
    return redirect(url_for('assignment_detail', assignment_id=assignment.id))

@app.route('/submission/<int:submission_id>/download')
@login_required
def download_submission(submission_id):
    submission = Submission.query.get_or_404(submission_id)
    assignment = submission.tarea
    
    # Solo el maestro o el estudiante que entregó pueden descargar
    if session['user_id'] != assignment.teacher_id and session['user_id'] != submission.student_id:
        flash('No tienes permiso para descargar esta entrega.', 'danger')
        return redirect(url_for('dashboard_student'))
    
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], submission.url_archivo)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    
    flash('Archivo no encontrado.', 'danger')
    return redirect(url_for('assignment_detail', assignment_id=assignment.id))

# ============== MANEJO DE ERRORES ==============

@app.errorhandler(404)
def not_found(error):
    return render_template('404.html'), 404

@app.errorhandler(500)
def server_error(error):
    return render_template('500.html'), 500

# ============== CREAR TABLAS ==============

with app.app_context():
    db.create_all()

    # migración ligera para bases existentes
    assignment_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(assignment)")).fetchall()]
    if 'target_student_id' not in assignment_columns:
        db.session.execute(text("ALTER TABLE assignment ADD COLUMN target_student_id INTEGER"))

    material_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(material)")).fetchall()]
    if 'target_student_id' not in material_columns:
        db.session.execute(text("ALTER TABLE material ADD COLUMN target_student_id INTEGER"))
    if 'board_status' not in material_columns:
        db.session.execute(text("ALTER TABLE material ADD COLUMN board_status VARCHAR(20) DEFAULT 'published'"))
    if 'scheduled_for' not in material_columns:
        db.session.execute(text("ALTER TABLE material ADD COLUMN scheduled_for DATETIME"))
    if 'planning_items' not in material_columns:
        db.session.execute(text("ALTER TABLE material ADD COLUMN planning_items TEXT"))
    if 'word_content' not in material_columns:
        db.session.execute(text("ALTER TABLE material ADD COLUMN word_content TEXT"))

    user_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(user)")).fetchall()]
    if 'notification_sound_enabled' not in user_columns:
        db.session.execute(text("ALTER TABLE user ADD COLUMN notification_sound_enabled BOOLEAN DEFAULT 1"))
    if 'paypal_account_email' not in user_columns:
        db.session.execute(text("ALTER TABLE user ADD COLUMN paypal_account_email VARCHAR(160)"))
    if 'paypal_account_name' not in user_columns:
        db.session.execute(text("ALTER TABLE user ADD COLUMN paypal_account_name VARCHAR(160)"))
    if 'nationality' not in user_columns:
        db.session.execute(text("ALTER TABLE user ADD COLUMN nationality VARCHAR(80)"))
    if 'preferred_payment_method' not in user_columns:
        db.session.execute(text("ALTER TABLE user ADD COLUMN preferred_payment_method VARCHAR(30)"))
    if 'preferred_bank_country' not in user_columns:
        db.session.execute(text("ALTER TABLE user ADD COLUMN preferred_bank_country VARCHAR(80)"))
    if 'paypal_data_opt_in' not in user_columns:
        db.session.execute(text("ALTER TABLE user ADD COLUMN paypal_data_opt_in BOOLEAN DEFAULT 0"))
    if 'session_nonce' not in user_columns:
        db.session.execute(text("ALTER TABLE user ADD COLUMN session_nonce INTEGER DEFAULT 0"))
    db.session.execute(text("UPDATE user SET notification_sound_enabled = 1 WHERE notification_sound_enabled IS NULL"))
    db.session.execute(text("UPDATE user SET paypal_data_opt_in = 0 WHERE paypal_data_opt_in IS NULL"))
    db.session.execute(text("UPDATE user SET session_nonce = 0 WHERE session_nonce IS NULL"))

    important_date_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(important_date)")).fetchall()]
    if 'meeting_link' not in important_date_columns:
        db.session.execute(text("ALTER TABLE important_date ADD COLUMN meeting_link VARCHAR(500)"))
    if 'emoji' not in important_date_columns:
        db.session.execute(text("ALTER TABLE important_date ADD COLUMN emoji VARCHAR(20)"))
    if 'reminder_note' not in important_date_columns:
        db.session.execute(text("ALTER TABLE important_date ADD COLUMN reminder_note VARCHAR(255)"))

    material_annotation_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(material_annotation)")).fetchall()]
    if 'page_number' not in material_annotation_columns:
        db.session.execute(text("ALTER TABLE material_annotation ADD COLUMN page_number INTEGER"))

    course_page_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(course_page)")).fetchall()]
    if 'target_student_id' not in course_page_columns:
        db.session.execute(text("ALTER TABLE course_page ADD COLUMN target_student_id INTEGER"))
    if 'is_published' not in course_page_columns:
        db.session.execute(text("ALTER TABLE course_page ADD COLUMN is_published BOOLEAN DEFAULT 1"))
    db.session.execute(text("UPDATE course_page SET is_published = 1 WHERE is_published IS NULL"))

    student_payment_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(student_payment)")).fetchall()]
    if 'paypal_account_email' not in student_payment_columns:
        db.session.execute(text("ALTER TABLE student_payment ADD COLUMN paypal_account_email VARCHAR(160)"))
    if 'paypal_account_name' not in student_payment_columns:
        db.session.execute(text("ALTER TABLE student_payment ADD COLUMN paypal_account_name VARCHAR(160)"))

    login_event_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(login_event)")).fetchall()]
    if not login_event_columns:
        LoginEvent.__table__.create(bind=db.engine, checkfirst=True)

    course_link_columns = [row[1] for row in db.session.execute(text("PRAGMA table_info(course_link)")).fetchall()]
    if not course_link_columns:
        CourseLink.__table__.create(bind=db.engine, checkfirst=True)

    db.session.commit()

    fixed_admin = User.query.filter(
        (User.email == FIXED_ADMIN_EMAIL) | (User.username == FIXED_ADMIN_USERNAME)
    ).first()

    if not fixed_admin:
        fixed_admin = User(
            username=FIXED_ADMIN_USERNAME,
            email=FIXED_ADMIN_EMAIL,
            password=generate_password_hash(FIXED_ADMIN_PASSWORD, method='pbkdf2:sha256'),
            role='teacher',
            nombre=FIXED_ADMIN_NAME
        )
        db.session.add(fixed_admin)
    else:
        fixed_admin.username = FIXED_ADMIN_USERNAME
        fixed_admin.email = FIXED_ADMIN_EMAIL
        fixed_admin.nombre = FIXED_ADMIN_NAME
        fixed_admin.role = 'teacher'
        fixed_admin.password = generate_password_hash(FIXED_ADMIN_PASSWORD, method='pbkdf2:sha256')

    db.session.execute(text("UPDATE user SET session_nonce = COALESCE(session_nonce, 0) + 1"))

    # limpiar usuarios de prueba para entorno de entrega
    db.session.execute(
        text(
            """
            DELETE FROM user
            WHERE LOWER(COALESCE(email, '')) IN ('estudiante@example.com', 'maestro@example.com')
              AND LOWER(COALESCE(email, '')) <> :fixed_email
            """
        ),
        {'fixed_email': FIXED_ADMIN_EMAIL.lower()}
    )

    db.session.execute(
        text(
            """
            DELETE FROM user
            WHERE LOWER(COALESCE(username, '')) IN ('estudiante', 'maestro')
              AND LOWER(COALESCE(email, '')) <> :fixed_email
            """
        ),
        {'fixed_email': FIXED_ADMIN_EMAIL.lower()}
    )

    db.session.commit()

# ============== EJECUTAR ==============

if __name__ == '__main__':
    # use port 5050 in case other typical ports are occupied
    app.run(debug=True, host='0.0.0.0', port=5050)
